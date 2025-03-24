# File Name: Minutes.py
# Authors: Javier Chung, Andy Dai, Antonio Lio, Jason Stuckless
# Description: This file handles generating the meeting minutes document

from docx import Document
import Transcript
import GPT
import os
import re
import docx2txt


def generateMinutes(meetingNotesList, filepath, minutesFilename, model_index=None):
    """
    Generate a meeting minutes document using the provided meetingNotesList

    Args:
        meetingNotesList: List of meeting notes sections
        filepath: Original transcript file path
        minutesFilename: Output filename for minutes document
        model_index: Index of the model used (0=DeepSeek R1, 1=DeepSeek API, 2=OpenAI API)
    """
    # Create MeetingNotes directory if it doesn't exist
    os.makedirs("MeetingNotes", exist_ok=True)

    # Ensure the file is saved to the MeetingNotes folder
    if not minutesFilename.startswith("MeetingNotes/"):
        minutesFilename = os.path.join("MeetingNotes", os.path.basename(minutesFilename))

    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading('Meeting Minutes', 0)

    # Define section titles
    section_titles = [
        "Opening Information",
        "Present Members",
        "Absent Members",
        "Agenda Approval",
        "Previous Meeting Minutes Approval",
        "Summary of Last Meeting",
        "Meeting Summary",
        "Actionable Items",
        "Adjournment"
    ]

    # Process sections
    for i, title in enumerate(section_titles):
        # Add section heading
        doc.add_heading(title, level=1)

        # Skip if we don't have enough sections
        if i >= len(meetingNotesList):
            continue

        content = meetingNotesList[i]

        # Clean up specific numbered headers that are causing issues
        if i == 6:  # Meeting Summary
            # Remove "7. Detailed Summary and Key Points of the Topic of the Current Meeting"
            content = re.sub(r'^7\.\s+Detailed Summary and Key Points of the Topic of (?:the )?Current Meeting\s*$',
                             '', content, flags=re.MULTILINE)
        elif i == 7:  # Actionable Items
            # Remove "8. Action Items and Assigned Responsibilities"
            content = re.sub(r'^8\.\s+Action Items and Assigned Responsibilities\s*$',
                             '', content, flags=re.MULTILINE)
        elif i == 8:  # Adjournment
            # Remove "9. Adjournment Time"
            content = re.sub(r'^9\.\s+Adjournment Time\s*$',
                             '', content, flags=re.MULTILINE)

        # More general cleaning for any section
        # Remove any line that starts with a number followed by a period and matches a section title
        for section_title in section_titles:
            pattern = fr'^\d+\.\s+{re.escape(section_title)}.*$'
            content = re.sub(pattern, '', content, flags=re.MULTILINE | re.IGNORECASE)

        # Remove any line that ONLY has a number and a period followed by any text
        content = re.sub(r'^\d+\.\s+.*$', '', content, flags=re.MULTILINE)

        # Process each paragraph in the content
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Skip lines that only contain a section number (e.g., "7.")
                if re.match(r'^\d+\.$', paragraph):
                    continue

                # Skip duplicate section titles that might appear in the content
                if any(section.lower() in paragraph.lower() for section in section_titles):
                    continue

                # Add the paragraph content
                doc.add_paragraph(paragraph)

    # Save the document
    doc.save(minutesFilename)
    print(f"Meeting minutes document saved as: {minutesFilename}")
    return minutesFilename