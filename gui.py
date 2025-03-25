# File Name: gui.py
# Authors: Javier Chung, Andy Dai, Antonio Lio, Jason Stuckless
# Date: Oct 26, 2024
# Description: This is the main page containing the GUI

# pip install PyQt5
# pip install python-docx       # For working with docx files
# pip install Pillow            # For image handling
from PyQt5 import QtCore
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QStackedWidget, QPushButton, QVBoxLayout, QLineEdit, QLabel, QHBoxLayout, QSizePolicy, QGraphicsDropShadowEffect, QGridLayout, QFrame, QScrollArea, QComboBox, QSpacerItem
from PyQt5.QtGui import QPalette, QColor, QFont, QPixmap, QCursor
from PyQt5.QtCore import Qt, pyqtSignal, QThread
from PyQt5.QtWidgets import QMessageBox, QFileDialog, QButtonGroup
import sys
from FileSystem import FileSystem
import pam
import openai
import speakerDiarization
import Recorder
from login import LoginSystem
from docx import Document
import os
import io
from PIL import Image, ImageQt

# connecting to openai
openai.api_key = os.getenv(
    "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")


# setting up global values
documentFont = QFont("Times New Roman", 12)  # Default font and font size

# Global shadow
shadow = QGraphicsDropShadowEffect()
shadow.setBlurRadius(40)
shadow.setXOffset(10)
shadow.setYOffset(10)
shadow.setColor(QColor(0, 0, 0, 100))

 # -------------------------------------- LOGIN PAGE ----------------------------------------- #

# this gui allows users to interact and login to their accounts
class LoginWindow(QWidget):
    def __init__(self, stack):
        super().__init__()

        self.stack = stack
        self.login_system = LoginSystem()  # Initialize login system
        self.is_login_mode = True  # Flag to toggle between login and register mode

        # Setting the background color of the page
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#DFDFDF'))
        self.setPalette(palette)
        self.setAutoFillBackground(True)
        self.setStyleSheet("QLabel{color: #000000;}")  # setting the default text color for all labels

        # Adding content onto the gui page
        mainLayout = QHBoxLayout()  # setting main layout to be horizontal

        # Setting up the layout on the left side
        layoutLeft = QVBoxLayout()
        logoImageLabel = QLabel(self)

        # Loading the image
        logo = QPixmap("./images/pamLogo.png")
        logoImageLabel.setPixmap(logo)

        # setting up the size
        logoImageLabel.setScaledContents(True)
        logoImageLabel.setFixedHeight(375)
        logoImageLabel.setFixedWidth(375)
        layoutLeft.addWidget(logoImageLabel)
        layoutLeft.setAlignment(Qt.AlignCenter)

        # setting up the right widget to contain the color
        rightWidget = QWidget()
        rightWidget.setStyleSheet(
            "background-color: #ffffff; border-radius: 10px; font-family: Arial; text-align: center;")

        rightWidget.setGraphicsEffect(shadow)

        # Setting up the layout on the right side to be vertical
        self.layoutRight = QVBoxLayout(rightWidget)
        self.layoutRight.setContentsMargins(50, 25, 50, 25)
        self.layoutRight.setSpacing(20)

        # Create title label for login/register
        self.loginLabel = QLabel("LOGIN")
        loginFont = QFont("Arial", 40, QFont.Bold)
        self.loginLabel.setFont(loginFont)
        self.layoutRight.addWidget(self.loginLabel)

        # Create input widgets
        # Setting up input on the GUI
        self.inputUsername = QLineEdit()  # Variable can be used anywhere in class when self is in the front
        self.inputPassword = QLineEdit()

        # Email field for registration
        self.inputEmail = QLineEdit()
        self.inputEmail.setVisible(False)  # Hidden by default in login mode

        # styling for the inputs
        inputStyle = "background-color: #E8E8E8; color: #7A7A7A; padding: 10px; font-size: 20px;"
        self.inputUsername.setStyleSheet(inputStyle)
        self.inputPassword.setStyleSheet(inputStyle)
        self.inputEmail.setStyleSheet(inputStyle)

        # Setting up placeholder text
        self.inputUsername.setPlaceholderText("Username")
        self.inputPassword.setPlaceholderText("Password")
        self.inputPassword.setEchoMode(QLineEdit.Password)  # Hide password
        self.inputEmail.setPlaceholderText("Email (optional)")

        # adding input for the username and password to the layout
        self.layoutRight.addWidget(self.inputUsername)
        self.layoutRight.addWidget(self.inputPassword)
        self.layoutRight.addWidget(self.inputEmail)  # Add email field but it's hidden initially

        # creating buttons for login and register
        self.loginButton = QPushButton(text="Login", parent=self)
        self.loginButton.setStyleSheet("""
            QPushButton {
                background-color: #E9E9E9;
                font-size: 15px;
                color: #000000;
                padding: 10px;
                border-radius: 8px;
            }
            QPushButton:hover {
                background-color: #D6D6D6;
            }
        """)
        self.loginButton.setFixedWidth(150)
        self.loginButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
        self.loginButton.clicked.connect(self.handle_auth)
        self.layoutRight.addWidget(self.loginButton)

        # Add "Switch to Register" button
        self.switchModeButton = QPushButton(text="Create Account", parent=self)
        self.switchModeButton.setStyleSheet(
            "background-color: transparent; font-size: 15px; color: #0000FF; padding: 5px; border: none;")
        self.switchModeButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
        self.switchModeButton.clicked.connect(self.toggle_mode)
        self.layoutRight.addWidget(self.switchModeButton)

        # centering elements in the layout
        self.layoutRight.setAlignment(Qt.AlignCenter)
        self.layoutRight.setAlignment(self.loginLabel, Qt.AlignCenter)
        self.layoutRight.setAlignment(self.loginButton, Qt.AlignCenter)
        self.layoutRight.setAlignment(self.switchModeButton, Qt.AlignCenter)

        # setting up how the layouts will be constructed and the width of them
        mainLayout.addLayout(layoutLeft)
        mainLayout.addWidget(rightWidget)

        # first parameter is the layout (added in a mainLayout in order and second is the strech)
        mainLayout.setStretch(0, 1)  # Left layout
        mainLayout.setStretch(1, 1)  # Right layout

        self.setLayout(mainLayout)

    def toggle_mode(self):
        """Toggle between login and register modes"""
        try:
            self.is_login_mode = not self.is_login_mode

            # Update UI based on mode
            if self.is_login_mode:
                self.loginLabel.setText("LOGIN")
                self.loginButton.setText("Login")
                self.switchModeButton.setText("Create Account")
                self.inputEmail.setVisible(False)
            else:
                self.loginLabel.setText("REGISTER")
                self.loginButton.setText("Register")
                self.switchModeButton.setText("Back to Login")
                self.inputEmail.setVisible(True)

            # Clear fields
            self.inputUsername.clear()
            self.inputPassword.clear()
            self.inputEmail.clear()
        except Exception as e:
            print(f"Error toggling mode: {e}")

    def handle_auth(self):
        """Handle both login and registration based on current mode"""
        try:
            username = self.inputUsername.text()
            password = self.inputPassword.text()

            # Basic validation
            if not username or not password:
                self.login_system.show_message(
                    self,
                    "Error",
                    "Username and password are required",
                    QMessageBox.Warning
                )
                return

            if self.is_login_mode:
                # Handle login
                success, message = self.login_system.authenticate_user(username, password)
                if success:
                    # go to next page
                    self.stack.setCurrentIndex(1)
                else:
                    self.login_system.show_message(
                        self,
                        "Login Failed",
                        message,
                        QMessageBox.Warning
                    )
            else:
                # Handle registration
                email = self.inputEmail.text()
                success, message = self.login_system.register_user(username, password, email)
                if success:
                    self.login_system.show_message(
                        self,
                        "Registration Successful",
                        "Your account has been created. You can now log in."
                    )
                    # Switch back to login mode
                    self.toggle_mode()
                else:
                    self.login_system.show_message(
                        self,
                        "Registration Failed",
                        message,
                        QMessageBox.Warning
                    )
        except Exception as e:
            print(f"Authentication error: {e}")
            self.login_system.show_message(
                self,
                "Error",
                "An unexpected error occurred",
                QMessageBox.Critical
            )

    def login(self):
        """For backward compatibility with existing code"""
        self.handle_auth()

# -------------------------------------- MAIN PAGE ----------------------------------------- #

# this contains the GUI for the main page of our application
class MainPage(QWidget):
    def __init__(self, stack):
        super().__init__()

        self.stack = stack
        self.fileSystem = FileSystem()  # initializing the FileSystem
        self.currentView = "minutes"  # Track which view we're showing: "minutes", "transcripts", or "recordings"

        # setting the background color
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#DFDFDF'))
        self.setPalette(palette)
        self.setAutoFillBackground(True)
        self.setStyleSheet("QLabel{color: #000000;}")  # setting the default text color for all labels

        # setting up the layout
        mainLayout = QVBoxLayout()

        # Creating the header bar at the top which will include the logo and buttons
        navBarWidget = QWidget()
        navBarWidget.setStyleSheet(
            "background-color: #ffffff; border-radius: 10px; font-family: Arial; text-align: center;")
        navBarWidget.setFixedHeight(150)
        navBarLayout = QHBoxLayout(navBarWidget)
        navBarLayout.setContentsMargins(50, 25, 50, 25)

        # Adding the logo
        logoImageLabel = QLabel(self)

        # Loading the image
        logo = QPixmap("./images/pamLogo (1).png")
        logoImageLabel.setPixmap(logo)

        # setting up the size
        logoImageLabel.setScaledContents(True)
        logoImageLabel.setFixedHeight(130)
        logoImageLabel.setFixedWidth(175)

        navBarLayout.addWidget(logoImageLabel)

        def createImageButton(image_path, text, click_action):
            """ Creates a QPushButton with an icon above text while keeping a solid background. """
            button = QPushButton()
            button.setFixedSize(150, 100)  # Set button size
            button.setCursor(QCursor(QtCore.Qt.PointingHandCursor))

            layout = QVBoxLayout(button)
            layout.setSpacing(5)
            layout.setAlignment(QtCore.Qt.AlignCenter)

            iconLabel = QLabel()
            iconPixmap = QPixmap(image_path).scaled(40, 40, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
            iconLabel.setPixmap(iconPixmap)
            iconLabel.setAlignment(QtCore.Qt.AlignCenter)

            textLabel = QLabel(text)
            textLabel.setAlignment(QtCore.Qt.AlignCenter)
            textLabel.setStyleSheet("font-size: 15px; color: #000000;")

            layout.addWidget(iconLabel)
            layout.addWidget(textLabel)
            button.setLayout(layout)

            button.setStyleSheet("""
                QPushButton {
                    background-color: #ffffff;
                    border-radius: 8px;
                    border: none;
                }
                QPushButton:hover {
                    background-color: #E9E9E9;
                }
            """)

            button.clicked.connect(click_action)
            return button

        # Create buttons with images
        summarizeButton = createImageButton("./images/summarize_icon.png", "Summarize", self.navigateSummarize)
        recordButton = createImageButton("./images/record_icon.png", "Record Meeting", self.navigateRecordAudio)

        # Add buttons to navbar
        navBarLayout.addWidget(summarizeButton)
        navBarLayout.addWidget(recordButton)

        # Create toggle view control
        toggleViewWidget = self.createToggleView()

        # Create a scrollable area for the document grid
        scrollArea = QScrollArea()
        scrollArea.setWidgetResizable(True)
        scrollArea.setFrameShape(QFrame.NoFrame)
        scrollArea.setStyleSheet("""
            QScrollArea {
                border: none;
                border-radius: 10px;
                background-color: #ffffff;
            }
            QScrollArea > QWidget {
                border-radius: 10px;
                background-color: #ffffff;
            }
            /* Scrollbar Styling */
            QScrollBar:vertical {
                border: none;
                background: transparent;
                width: 15px;
                margin: 2px 5px 2px 0px;
            }
            QScrollBar::handle:vertical {
                background: rgba(0, 0, 0, 0.2);
                border-radius: 5px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                background: none;
                border: none;
            }
        """)

        # Create a content widget to hold the documents
        contentWidget = QWidget()
        contentWidget.setStyleSheet("""
            background-color: #ffffff;
            border-radius: 10px;
            font-family: Arial;
            color: #000000;
            padding: 20px;
        """)

        # Set up the grid layout inside the content widget
        self.documentsLayout = QGridLayout(contentWidget)
        self.documentsLayout.setContentsMargins(50, 25, 50, 25)
        self.documentsLayout.setSpacing(20)

        # Set the content widget inside the scroll area
        scrollArea.setWidget(contentWidget)

        # Add the navbar and toggle view to the main layout
        mainLayout.addWidget(navBarWidget)
        mainLayout.addWidget(toggleViewWidget)
        mainLayout.addWidget(scrollArea)

        # Set the main layout for the window
        self.setLayout(mainLayout)

        # Update the grid layout for documents
        self.updateDocuments(view_type="minutes")

    def createToggleView(self):
        """Creates a toggle control to switch between transcripts, recordings, and minutes view"""
        # Create a container for the toggle
        toggleContainer = QWidget()
        toggleContainer.setStyleSheet("""
            background-color: #ffffff;
            border-radius: 10px;
            padding: 10px;
        """)
        toggleContainer.setFixedHeight(60)

        # Create horizontal layout for the toggle
        toggleLayout = QHBoxLayout(toggleContainer)
        toggleLayout.setContentsMargins(20, 5, 20, 5)

        # Create label for the toggle
        viewLabel = QLabel("View:")
        viewLabel.setStyleSheet("font-size: 16px; font-weight: bold;")
        toggleLayout.addWidget(viewLabel)

        # Add spacer
        toggleLayout.addSpacerItem(QSpacerItem(20, 20, QSizePolicy.Fixed, QSizePolicy.Minimum))

        # Create a button group for mutual exclusion
        self.viewButtonGroup = QButtonGroup(self)

        # Create buttons for toggle options with consistent styling
        buttonStyle = """
            QPushButton {
                background-color: #E9E9E9;
                border-radius: 8px;
                padding: 8px 15px;
                font-size: 14px;
            }
            QPushButton:checked {
                background-color: #4682B4;
                color: white;
            }
            QPushButton:hover {
                background-color: #D6D6D6;
            }
        """

        self.minutesRadio = QPushButton("Minutes")
        self.minutesRadio.setCheckable(True)
        self.minutesRadio.setChecked(True)  # Default to minutes view
        self.minutesRadio.setStyleSheet(buttonStyle)

        self.transcriptsRadio = QPushButton("Transcripts")
        self.transcriptsRadio.setCheckable(True)
        self.transcriptsRadio.setStyleSheet(buttonStyle)

        self.recordingsRadio = QPushButton("Recordings")
        self.recordingsRadio.setCheckable(True)
        self.recordingsRadio.setStyleSheet(buttonStyle)

        # Add buttons to the button group
        self.viewButtonGroup.addButton(self.minutesRadio)
        self.viewButtonGroup.addButton(self.transcriptsRadio)
        self.viewButtonGroup.addButton(self.recordingsRadio)

        # Add buttons to layout
        toggleLayout.addWidget(self.minutesRadio)
        toggleLayout.addWidget(self.transcriptsRadio)
        toggleLayout.addWidget(self.recordingsRadio)

        # Add spacer to push buttons to the left
        toggleLayout.addStretch(1)

        # Connect toggle buttons to update view
        self.minutesRadio.clicked.connect(lambda: self.updateDocuments(view_type="minutes"))
        self.transcriptsRadio.clicked.connect(lambda: self.updateDocuments(view_type="transcripts"))
        self.recordingsRadio.clicked.connect(lambda: self.updateDocuments(view_type="recordings"))

        return toggleContainer


    # go to summarize page
    def navigateSummarize(self):
        sumPage = self.stack.widget(2)
        sumPage.resetFileInput() # rest the summarization page
        self.stack.setCurrentIndex(2)

    # go to record audio page
    def navigateRecordAudio(self):
        recordAudioPage = self.stack.widget(3)
        self.stack.setCurrentIndex(3)

    # go to document page
    def navigateDocument(self, fileName):
        # Send the file name to the DocumentPage
        self.docPage = DocumentPage(self.stack, fileName)
        
        # Adding the DocumentPage to the stack and switching to it
        self.stack.addWidget(self.docPage)
        self.stack.setCurrentWidget(self.docPage)

    def updateDocuments(self, view_type=None):
        """Update the documents grid based on the selected view type"""
        # If view_type is None, use current state
        if view_type is None:
            view_type = self.currentView
        else:
            self.currentView = view_type

        # Clear existing widgets from the grid first to avoid duplicates
        for i in reversed(range(self.documentsLayout.count())):
            widget = self.documentsLayout.itemAt(i).widget()
            if widget:
                widget.deleteLater()

        # Retrieve files based on the view type
        if view_type == "recordings":
            self.fileSystem.searchAudioDirectory()
            fileNames = self.fileSystem.audioFileNames
            iconPath = "./images/audioIcon.png"  # You need to create/add this icon
            directory = self.fileSystem.audioDirectory
            getSize = self.fileSystem.getAudioSize
        elif view_type == "transcripts":
            self.fileSystem.searchTranscriptDirectory()
            fileNames = self.fileSystem.transcriptFileNames
            iconPath = "./images/docIcon.png"  # Reuse document icon or create a specific transcript icon
            directory = self.fileSystem.transcriptDirectory
            getSize = self.fileSystem.getTranscriptSize
        else:  # "minutes" is the default
            self.fileSystem.searchDirectory()
            fileNames = self.fileSystem.fileNames
            iconPath = "./images/docIcon.png"
            directory = self.fileSystem.fileDirectory
            getSize = self.fileSystem.getSize

        fileAmt = len(fileNames)  # Get the actual number of files
        columns = 5  # Keep exactly 5 columns at all times

        # Ensure all 5 columns stretch evenly
        for i in range(columns):
            self.documentsLayout.setColumnStretch(i, 1)

        row = 0
        documentWidgets = []  # Store widgets to resize them uniformly later

        while fileAmt > 0:
            for col in range(columns):
                if fileAmt > 0:
                    fileIndex = (row * columns) + col
                    if fileIndex >= len(fileNames):
                        break

                    fileName = fileNames[fileIndex]

                    # Grey box (container for the file button)
                    documentGridWidget = QWidget()
                    documentGridWidget.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.Preferred)
                    documentGridWidget.setMinimumWidth(150)
                    documentGridWidget.setMaximumWidth(300)
                    documentGridWidget.setMinimumHeight(200)  # Increased height
                    documentGridWidget.setStyleSheet("""
                        QWidget {
                            background-color: #868686; 
                            border-radius: 10px; 
                            font-family: Arial; 
                            color: #000000;
                        }
                    """)

                    # Apply click event for the entire grey box
                    documentGridWidget.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
                    if view_type == "recordings":
                        documentGridWidget.mousePressEvent = lambda event, name=fileName, dir=directory: self.openAudioFile(name, dir)
                    elif view_type == "transcripts":
                        # For transcripts, we pass the fileName
                        documentGridWidget.mousePressEvent = lambda event, name=fileName: self.navigateTranscript(name)
                    else:
                        documentGridWidget.mousePressEvent = lambda event, name=fileName: self.navigateDocument(name)

                    # Create layout for each document cell
                    documentGrid = QVBoxLayout(documentGridWidget)
                    documentGrid.setContentsMargins(10, 10, 10, 10)
                    documentGrid.setSpacing(10)  # Increased spacing

                    # Load the appropriate icon
                    docIconLabel = QLabel(documentGridWidget)

                    # Check if the icon exists, if not use a default
                    if not os.path.exists(iconPath):
                        # Use a fallback icon if the icon doesn't exist
                        iconPath = "./images/docIcon.png"

                    docIconPixmap = QPixmap(iconPath).scaled(70, 70, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    docIconLabel.setPixmap(docIconPixmap)
                    docIconLabel.setAlignment(Qt.AlignCenter)
                    docIconLabel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

                    # Create text label with explicit text elision control
                    fileNameLabel = QLabel(documentGridWidget)

                    # If filename is too long (>15 chars), wrap it to multiple lines
                    displayText = fileName
                    if len(fileName) > 20:
                        # Insert line breaks at appropriate points
                        words = fileName.split('_')
                        if len(words) > 1:
                            # If filename contains underscores, break at underscores
                            displayText = '\n'.join(words)
                        else:
                            # Otherwise, break at reasonable points
                            displayText = fileName[:15] + '\n' + fileName[15:]

                    fileNameLabel.setText(displayText)
                    fileNameLabel.setAlignment(Qt.AlignCenter)
                    fileNameLabel.setWordWrap(True)
                    fileNameLabel.setMinimumHeight(80)  # Ensure enough space for multiple lines
                    fileNameLabel.setStyleSheet("""
                        font-size: 14px; 
                        font-weight: bold;
                        color: black;
                        background-color: transparent;
                        padding: 5px;
                    """)

                    # Add elements to layout
                    documentGrid.addWidget(docIconLabel, alignment=Qt.AlignCenter)
                    documentGrid.addWidget(fileNameLabel, alignment=Qt.AlignCenter)
                    documentGrid.setStretch(0, 1)  # Icon gets less space
                    documentGrid.setStretch(1, 2)  # Text gets more space

                    # Add widget to the grid layout
                    self.documentsLayout.addWidget(documentGridWidget, row, col)

                    # Store reference for uniform resizing
                    documentWidgets.append(documentGridWidget)

                    fileAmt -= 1
                else:
                    break
            row += 1

        # Set more spacing between grid cells
        self.documentsLayout.setHorizontalSpacing(20)
        self.documentsLayout.setVerticalSpacing(20)

        # Force a UI refresh
        self.documentsLayout.update()
        self.updateGeometry()

    def navigateTranscript(self, fileName):
        """Navigate to a transcript file using the specialized TranscriptPage"""
        # Create a TranscriptPage instead of a DocumentPage
        self.transcriptPage = TranscriptPage(self.stack, fileName)

        # Add the page to the stack and switch to it
        self.stack.addWidget(self.transcriptPage)
        self.stack.setCurrentWidget(self.transcriptPage)

    def openAudioFile(self, fileName, directory):
        """Opens an audio file using the system's default audio player"""
        import os
        import subprocess
        import platform

        # Get the full path to the audio file
        filePath = os.path.join(directory, fileName)

        # Check if the file exists
        if not os.path.exists(filePath):
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Warning)
            msg.setText(f"File not found: {filePath}")
            msg.setWindowTitle("File Not Found")
            msg.exec_()
            return

        # Open the audio file with the system's default program
        try:
            system = platform.system()

            if system == "Windows":
                os.startfile(filePath)
            elif system == "Darwin":  # macOS
                subprocess.call(["open", filePath])
            else:  # Linux and variants
                subprocess.call(["xdg-open", filePath])

            print(f"Opening audio file: {filePath}")
        except Exception as e:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Warning)
            msg.setText(f"Error opening file: {str(e)}")
            msg.setWindowTitle("Error")
            msg.exec_()


# ------------------- Record Audio Thread ------------------- #
class RecorderThread(QThread):
    # Defining the signals that will be used to send messages to and from the thread
    recordingFinished = pyqtSignal(str)
    stopRecording = pyqtSignal()
    def __init__(self):
        super().__init__()
        # Connecting the stopRecording signal to the stopRecording method
        self.stopRecording.connect(self.stopRecordingMethod)
        
    def run(self):

        # Recording the audio
        filename = Recorder.recordAudio()

        # Emit signal when recording is finished
        self.recordingFinished.emit(filename)

    def stopRecordingMethod(self):
        print("Recording stopped")
        Recorder.setStopRecording(True)
        
# -------------------------------------- RECORD AUDIO PAGE ----------------------------------------- #
# this contains the GUI for the main page of our application
class RecordAudioPage(QWidget):
    record = False;

    def __init__(self, stack):

        def createImageButton(image_path, text, eventOnClick):
            """
            Creates a QPushButton with an icon above text while keeping a solid background.
            """
            # Create a QPushButton without text initially
            button = QPushButton()
            button.setFixedSize(300, 300) 

            # Tracking Button state to when there is a clicked event
            button.isRecording = False
            button.finsishedRecording = False
            self.recorderThread = None


            # Create a layout to stack icon and text
            layout = QVBoxLayout(button)
            layout.setSpacing(5)  # Space between icon and text
            layout.setAlignment(QtCore.Qt.AlignCenter)  # Center align content

            # Load the icon as a QLabel
            iconLabel = QLabel()
            iconPixmap = QPixmap(image_path).scaled(100, 100, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
            iconLabel.setPixmap(iconPixmap)
            iconLabel.setAlignment(QtCore.Qt.AlignCenter)

            # Create a QLabel for the text
            textLabel = QLabel(text)
            textLabel.setAlignment(QtCore.Qt.AlignCenter)
            textLabel.setStyleSheet("font-size: 30px; color: #000000;")

            # Add widgets to layout
            layout.addWidget(iconLabel)
            layout.addWidget(textLabel)

            # Apply the layout to the button
            button.setLayout(layout)

            # Apply button styles
            button.setStyleSheet("""
                QPushButton {
                    background-color: #ffffff;
                    border-radius: 8px;
                    border: 2px solid black;
                }
                QPushButton:hover {
                    background-color: #E9E9E9;
                }
            """)
        
        # This method will change the color of the button when it is pressed
            def toggleColor():
                if button.isRecording:
                    button.isRecording = False
                    button.setStyleSheet("""
                        QPushButton {
                            background-color: #ffffff;
                            border-radius: 8px;
                            border: 2px solid black;
                        }
                        QPushButton:hover {
                            background-color: #E9E9E9;
                        }
                    """)
                    button.finsishedRecording = True
                else:
                    button.isRecording = True

                    button.setStyleSheet("""
                        QPushButton {
                            background-color: red;
                            border-radius: 8px;
                            border: 2px solid black;
                        }
                    """)

                # Call the provided event function
                eventOnClick(button.isRecording, button.finsishedRecording)

                button.finsishedRecording = False

            # connected the button to the nest method to change the color which will then connect it to the main method of recording
            button.clicked.connect(toggleColor)

            return button

        super().__init__()

        self.stack = stack

        # setting the background color
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#DFDFDF'))
        self.setPalette(palette)
        self.setAutoFillBackground(True)

        # setting up the layout
        mainLayout = QVBoxLayout()

        # ----------------------------- Navbar START ----------------------------- #
        navBarWidget = QWidget()
        navBarWidget.setStyleSheet(
            "background-color: #ffffff; border-radius: 10px; font-family: Arial; text-align: center;")
        navBarWidget.setFixedHeight(100)

        # Set up the layout for the navbar
        navBarLayout = QHBoxLayout(navBarWidget)
        navBarLayout.setContentsMargins(0, 0, 0, 0)
        navBarLayout.setSpacing(0)

        # Back Button (Top Left)
        backButtonContainer = QVBoxLayout()
        backButtonContainer.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
        backButtonContainer.setContentsMargins(10, 10, 0, 0)

        # Back Button with Left Arrow Icon
        backButton = QPushButton(parent=self)
        backButton.setFixedSize(75, 75)
        backButton.setStyleSheet("""
            QPushButton {
                background-color: #E9E9E9; 
                font-size: 20px; 
                color: #000000; 
            }
            QPushButton:hover {
                background-color: #D6D6D6;
            }
        """)
        backButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
        backButton.setText("←")
        backButton.setGraphicsEffect(shadow)
        backButton.clicked.connect(self.navigateHome)

        backButtonContainer.addWidget(backButton)
        navBarLayout.addLayout(backButtonContainer)

        # Add left spacer
        navBarLayout.addStretch(1)

        # Add the logo (Centered)
        logoImageLabel = QLabel(self)
        logo = QPixmap("./images/pamLogo (1).png")
        logoImageLabel.setPixmap(logo)
        logoImageLabel.setScaledContents(True)
        logoImageLabel.setFixedHeight(125)
        logoImageLabel.setFixedWidth(150)
        navBarLayout.addWidget(logoImageLabel, alignment=QtCore.Qt.AlignCenter)

        # Add right spacer
        navBarLayout.addStretch(1)

        # Invisible Empty Widget to match Back Button size
        emptyWidget = QWidget()
        emptyWidget.setFixedSize(85, 85)
        navBarLayout.addWidget(emptyWidget)
        # ----------------------------- Navbar END ----------------------------- #

        # -------------------- Recording Layout where the user can select record ---------------------- #

        # Create a QWidget container for styling
        recordWidget = QWidget()
        recordWidget.setStyleSheet("""
            QWidget {
                background-color: #ffffff;  
                border-radius: 10px;        
                font-family: Arial;
                text-align: center;
                padding-top: 10px;
            }
        """)

        recordWidget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)


        # Creating a label with specified settings
        # style sheet - bold, size and color
        # static width and dynamic height for word wrapping
        self.instructionLabel = QLabel("   When Recording, Get Each Speaker to introduce themselves with the format of 'Hi, I am <First Name> <Last Name>' to register speaker names ")
        self.instructionLabel.setStyleSheet("""
            font-size: 16px;
            color: black;
            font-weight: bold;
            background-color: #FFFACD;  
            border: 2px solid #FFD700; 
            padding: 5px;  
        """)
        self.instructionLabel.setFixedWidth(self.width() // 2)
        self.instructionLabel.setWordWrap(True)  # Enable word wrapping
        self.instructionLabel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)

        # Creating the layout for the record button section
        recordLayout = QVBoxLayout(recordWidget)
        recordLayout.setAlignment(Qt.AlignCenter)

        # Create buttons with images
        self.recordButton = createImageButton("./images/record_icon.png", "Record Meeting", self.recordAudio)
        self.recordButton.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.drawDropDown()

        # Add the button to the layout

        recordLayout.addWidget(self.instructionLabel, alignment=Qt.AlignCenter)
        recordLayout.addSpacing(20)

        recordLayout.addWidget(self.dropdown, alignment=Qt.AlignCenter)
        recordLayout.addSpacing(10)

        recordLayout.addWidget(self.recordButton, alignment=Qt.AlignCenter)

        # Adding components to mainLayout
        mainLayout.addWidget(navBarWidget)
        mainLayout.addWidget(recordWidget)

        # Set the layout to the window
        self.setLayout(mainLayout)

    # Recording the audio through a thread
    def recordAudio(self, isRecording, finishedRecording):
        if not self.recorderThread or not self.recorderThread.isRunning():
            self.recorderThread = RecorderThread()

        # If the thread is finished executing, it will execute this method
        self.recorderThread.recordingFinished.connect(self.handleRecordingFinished)
        
        # Connect the stopRecording signal to a method that will stop the recording
        self.recorderThread.stopRecording.connect(self.stopRecording)

        # Starting the recording thread
        self.recorderThread.start()

        if (not isRecording and finishedRecording):
            print("send signal to stop thread")
            # Emit the stopRecording signal to stop the thread
            self.recorderThread.stopRecording.emit()
            Recorder.setStopRecording(True)

    def stopRecording(self):
        print("Recording is stopping now...")
        Recorder.setStopRecording(True)

    # this method will create the transcript and display a popup
    def handleRecordingFinished(self, recordedFile):
        # disconnecting the signal
        try:
            self.recorderThread.recordingFinished.disconnect(self.handleRecordingFinished)
            self.recorderThread.stopRecording.disconnect(self.stopRecording)
        except Exception as e:
            print("Signals already disconnected or error during disconnect:", e)

        # Pass the full path to the transcription function
        filepath = speakerDiarization.transcribeAndDiarize(self.dropdown.currentIndex(), recordedFile)

        # have a pop up saying transcript created in {folder}
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
        msg.setText(f"Transcript Successfully Created at filepath:{filepath}!\nRecording saved at: {recordedFile}")
        msg.setWindowTitle("Recording Finished")
        msg.exec_()

    # this function is used to draw a drop down bar
    def drawDropDown(self):
        self.dropdown = QComboBox()
        self.dropdown.addItems(["Medium en (Accurate)", "Base en (Moderate)", "Tiny en (Fast)"])  
        self.dropdown.setFixedWidth(self.width() // 2)
        self.dropdown.setStyleSheet("""
            QComboBox {
                background-color: #F5F5F5;
                color: black;
                font-size: 16px;
                border: 2px solid #CCCCCC;
                border-radius: 10px;
                padding: 5px;
            }
        """)

    # this function navigates back to the main page
    def navigateHome(self):

        # Clearing the recording if there was a recording in progress
        if self.recorderThread and self.recorderThread.isRunning():
            self.recorderThread.stopRecording.emit()

            # Waiting until the thread is finished
            Recorder.setStopRecording(True)            
            self.recorderThread.wait()

        # resetting the button so that it is not red and recording is set to false
        self.recordButton.isRecording = False
        self.recordButton.setStyleSheet("""
            QPushButton {
                background-color: #ffffff;
                border-radius: 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #E9E9E9;
            }
        """)

        self.stack.setCurrentIndex(1)

# -------------------------------------- SUMMARIZATION PAGE ----------------------------------------- #
# this contains the GUI for summarizing a note
class SummarizationPage(QWidget):
    def __init__(self, stack):
        super().__init__()

        self.stack = stack

        # setting the background color
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#DFDFDF'))
        self.setPalette(palette)
        self.setAutoFillBackground(True)

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(40)
        shadow.setXOffset(10)
        shadow.setYOffset(10)
        shadow.setColor(QColor(0, 0, 0, 100))

        # setting up the layout
        mainLayout = QVBoxLayout()

        # ----------------------------- Navbar START ----------------------------- #
        navBarWidget = QWidget()
        navBarWidget.setStyleSheet(
            "background-color: #ffffff; border-radius: 10px; font-family: Arial; text-align: center;")
        navBarWidget.setFixedHeight(100)

        # Set up the layout for the navbar
        navBarLayout = QHBoxLayout(navBarWidget)
        navBarLayout.setContentsMargins(0, 0, 0, 0)
        navBarLayout.setSpacing(0)

        # Back Button (Top Left)
        backButtonContainer = QVBoxLayout()
        backButtonContainer.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
        backButtonContainer.setContentsMargins(10, 10, 0, 0)

        # Back Button with Left Arrow Icon
        backButton = QPushButton(parent=self)
        backButton.setFixedSize(75, 75)
        backButton.setStyleSheet("""
            QPushButton {
                background-color: #E9E9E9; 
                font-size: 20px; 
                color: #000000; 
            }
            QPushButton:hover {
                background-color: #D6D6D6;
            }
        """)
        backButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
        backButton.setText("←")
        backButton.setGraphicsEffect(shadow)
        backButton.clicked.connect(self.navigateHome)

        backButtonContainer.addWidget(backButton)
        navBarLayout.addLayout(backButtonContainer)

        # Add left spacer
        navBarLayout.addStretch(1)

        # Add the logo (Centered)
        logoImageLabel = QLabel(self)
        logo = QPixmap("./images/pamLogo (1).png")
        logoImageLabel.setPixmap(logo)
        logoImageLabel.setScaledContents(True)
        logoImageLabel.setFixedHeight(125)
        logoImageLabel.setFixedWidth(150)
        navBarLayout.addWidget(logoImageLabel, alignment=QtCore.Qt.AlignCenter)

        # Add right spacer
        navBarLayout.addStretch(1)

        # Invisible Empty Widget to match Back Button size
        emptyWidget = QWidget()
        emptyWidget.setFixedSize(85, 85)
        navBarLayout.addWidget(emptyWidget)
        # ----------------------------- Navbar END ----------------------------- #

        self.resetFileInput()
        self.drawDropDown()

        # adding the components to the mainLayout
        mainLayout.addWidget(navBarWidget)

        # Add spacing between navbar and drop box
        mainLayout.addSpacing(15)

        mainLayout.addWidget(self.dropBox)

        # Add spacing between drop box and drop down
        mainLayout.addSpacing(15)

        mainLayout.addWidget(self.dropdown)

        # Add a bit of spacing before the buttons
        mainLayout.addSpacing(10)

        # Create a horizontal layout for the buttons
        buttonLayout = QHBoxLayout()
        buttonLayout.addWidget(self.browseButton)
        buttonLayout.addWidget(self.summarizeButton)

        # Add button layout to main layout
        mainLayout.addLayout(buttonLayout)

        # aligning the widgets
        mainLayout.setAlignment(self.dropBox, Qt.AlignCenter)
        mainLayout.setAlignment(self.dropdown, Qt.AlignCenter)
        buttonLayout.setAlignment(Qt.AlignCenter)

        # setting the layout to the window
        self.setLayout(mainLayout)

    # built in function for drop events
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()
    
    # built in function for drop events
    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    # This function is a built in function to react to a drop event
    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            fileName = event.mimeData().urls()[0].toLocalFile()
            event.setDropAction(Qt.CopyAction)
            self.dropBox.setText(fileName.split('/')[-1])

    # This function will invoke the summarization component
    def summarize(self,filename):
        pam.summarize(filename, self.dropdown.currentIndex())
        mainPage = self.stack.widget(1)
        mainPage.updateDocuments() # update the main page to reflect the files in the directory
        self.stack.setCurrentIndex(1)

    # this function navigates back to the main page
    def navigateHome(self):
        self.stack.setCurrentIndex(1)

    # this function is used to draw a drop down bar
    def drawDropDown(self):
        self.dropdown = QComboBox()
        self.dropdown.addItems(["DeepSeek 1.5B Local (slower/less accurate/free)", "DeepSeek API (faster/more accurate/cheap)", "OpenAI API (fastest/best/expensive)"])
        self.dropdown.setSizeAdjustPolicy(QComboBox.AdjustToContents)
        self.dropdown.setStyleSheet("""
            QComboBox {
                background-color: #F5F5F5;
                color: black;
                font-size: 16px;
                border: 2px solid #CCCCCC;
                border-radius: 10px;
                padding: 5px;
            }
        """)
        
    # this will reset the file input for the drop box
    def resetFileInput(self):
        # if the drop box has already been created than we reset the text
        if hasattr(self, 'dropBox'):
            self.dropBox.setText("Drop File Here")
        else:
            # creating a label where the user can drag and drop files to summarize
            self.setAcceptDrops(True)  # allows drags and drops
            self.dropBox = QLabel("Drop File Here", self)
            self.dropBox.setWordWrap(True)
            self.dropBox.setStyleSheet("border: 3px dashed #000000; font-size: 20px;")
            self.dropBox.setAlignment(Qt.AlignCenter)
            self.dropBox.setFixedHeight(300)
            self.dropBox.setFixedWidth(300)

            # Creating a button for summarizing with improved styling
            self.summarizeButton = QPushButton(text="Summarize", parent=self)
            self.summarizeButton.setStyleSheet("""
                QPushButton {
                    background-color: #E9E9E9; 
                    font-size: 15px; 
                    color: #000000; 
                    padding: 10px;
                    border-radius: 8px;
                }
                QPushButton:hover {
                    background-color: #D6D6D6;
                }
            """)
            self.summarizeButton.setFixedWidth(150)
            self.summarizeButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
            self.summarizeButton.setGraphicsEffect(shadow)
            self.summarizeButton.clicked.connect(lambda: self.summarize(self.dropBox.text()))

            # Creating a "Browse Files" button
            self.browseButton = QPushButton(text="Browse Files", parent=self)
            self.browseButton.setStyleSheet("""
                QPushButton {
                    background-color: #E9E9E9; 
                    font-size: 15px; 
                    color: #000000; 
                    padding: 10px;
                    border-radius: 8px;
                }
                QPushButton:hover {
                    background-color: #D6D6D6;
                }
            """)
            self.browseButton.setFixedWidth(150)
            self.browseButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
            self.browseButton.setGraphicsEffect(shadow)
            self.browseButton.clicked.connect(self.browseFiles)

    # Add this new method for file browsing functionality
    def browseFiles(self):
        options = QFileDialog.Options()
        filePath, _ = QFileDialog.getOpenFileName(
            self,
            "Select File to Summarize",
            "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;PDF Files (*.pdf)",
            options=options
        )

        if filePath:
            # Extract just the filename from the path
            fileName = filePath.split('/')[-1]
            # Update the drop box text with the selected file
            self.dropBox.setText(filePath)


# this contains the GUI for opening up a document
class DocumentPage(QWidget):
    def __init__(self, stack, fileName=None):
        super().__init__()

        self.stack = stack
        self.fileName = fileName
        self.fileSystem = FileSystem()  # initializing the file system

        # setting the background color
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#DFDFDF'))
        self.setPalette(palette)
        self.setAutoFillBackground(True)

        # setting up the layout
        mainLayout = QVBoxLayout()

        # ----------------------------- Navbar START ----------------------------- #
        navBarWidget = QWidget()
        navBarWidget.setStyleSheet(
            "background-color: #ffffff; border-radius: 10px; font-family: Arial; text-align: center;")
        navBarWidget.setFixedHeight(100)

        # Set up the layout for the navbar
        navBarLayout = QHBoxLayout(navBarWidget)
        navBarLayout.setContentsMargins(0, 0, 0, 0)
        navBarLayout.setSpacing(0)

        # Back Button (Top Left)
        backButtonContainer = QVBoxLayout()
        backButtonContainer.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
        backButtonContainer.setContentsMargins(10, 10, 0, 0)

        # Back Button with Left Arrow Icon
        backButton = QPushButton(parent=self)
        backButton.setFixedSize(75, 75)
        backButton.setStyleSheet("""
            QPushButton {
                background-color: #E9E9E9; 
                font-size: 20px; 
                color: #000000; 
            }
            QPushButton:hover {
                background-color: #D6D6D6;
            }
        """)
        backButton.setCursor(QCursor(QtCore.Qt.PointingHandCursor))
        backButton.setText("←")
        backButton.setGraphicsEffect(shadow)
        backButton.clicked.connect(self.navigateHome)

        backButtonContainer.addWidget(backButton)
        navBarLayout.addLayout(backButtonContainer)

        # Add left spacer
        navBarLayout.addStretch(1)

        # Add the logo (Centered)
        logoImageLabel = QLabel(self)
        logo = QPixmap("./images/pamLogo (1).png")
        logoImageLabel.setPixmap(logo)
        logoImageLabel.setScaledContents(True)
        logoImageLabel.setFixedHeight(125)
        logoImageLabel.setFixedWidth(150)
        navBarLayout.addWidget(logoImageLabel, alignment=QtCore.Qt.AlignCenter)

        # Add right spacer
        navBarLayout.addStretch(1)

        # Invisible Empty Widget to match Back Button size
        emptyWidget = QWidget()
        emptyWidget.setFixedSize(85, 85)
        navBarLayout.addWidget(emptyWidget)
        # ----------------------------- Navbar END ----------------------------- #

        # ----------------------------- File Specifications Bar ----------------------------- #
        fileSpecWidget = QWidget()
        fileSpecWidget.setStyleSheet("""
            background-color: #ffffff;
            border-radius: 10px;
            font-family: Arial;
            color: #000000;
            margin-top: 10px;
            margin-bottom: 10px;
            background-color: #D1D6D7;
        """)
        fileSpecWidget.setFixedHeight(100)

        fileSpecLayout = QHBoxLayout(fileSpecWidget)

        # File name
        fileNameLabel = QLabel(f"File: {fileName}")
        fileNameLabel.setFont(QFont("Times New Roman", 12))
        fileSpecLayout.addWidget(fileNameLabel)

        # Add spacer between file name and file size
        fileSpecLayout.addStretch(1)

        # File size
        try:
            fileSize = self.fileSystem.getSize(fileName)
            fileSizeLabel = QLabel(f"Size: {fileSize}")
            fileSizeLabel.setFont(QFont("Times New Roman", 12))
            fileSpecLayout.addWidget(fileSizeLabel)
        except Exception as e:
            print(f"Error getting file size: {str(e)}")

        # ----------------------------- Document Content ----------------------------- #

        # Creating a scroll area for document content
        scrollArea = QScrollArea()
        scrollArea.setWidgetResizable(True)
        scrollArea.setFrameShape(QFrame.NoFrame)
        scrollArea.setStyleSheet("""
            QScrollArea {
                border: none;
                border-radius: 10px;
                background-color: #ffffff;
            }
            QScrollArea > QWidget {
                border-radius: 10px;
                background-color: #ffffff;
            }
            /* Scrollbar Styling */
            QScrollBar:vertical {
                border: none;
                background: transparent;
                width: 10px;
                margin: 2px 5px 2px 0px;
            }
            QScrollBar::handle:vertical {
                background: rgba(0, 0, 0, 0.2);
                border-radius: 5px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                background: none;
                border: none;
            }
        """)

        # Create content widget for scroll area
        contentWidget = QWidget()
        contentWidget.setStyleSheet("""
            background-color: #ffffff;
            border-radius: 10px;
            font-family: Arial;
            color: #000000;
            padding: 20px;
        """)

        # Main document layout
        documentLayout = QVBoxLayout(contentWidget)
        documentLayout.setSpacing(15)

        try:
            # Get the full path to the document
            fullPath = os.path.join(self.fileSystem.fileDirectory, fileName)

            # Check if file exists
            if not os.path.exists(fullPath):
                errorLabel = QLabel(f"Error: File not found at {fullPath}")
                errorLabel.setAlignment(Qt.AlignCenter)
                documentLayout.addWidget(errorLabel)
            else:
                # Load the document using python-docx
                doc = Document(fullPath)

                # Add a horizontal line at the top of the document content
                line = QFrame()
                line.setFrameShape(QFrame.HLine)
                line.setFrameShadow(QFrame.Sunken)
                line.setStyleSheet("border: 1px solid #cccccc;")
                documentLayout.addWidget(line)

                # Add some spacing after the line
                documentLayout.addSpacing(10)

                # Process paragraphs
                for para in doc.paragraphs:
                    if not para.text.strip():  # Skip empty paragraphs
                        continue

                    # Create paragraph label
                    paraLabel = QLabel()

                    # Apply paragraph style
                    if para.style.name.startswith('Heading'):
                        # Get heading level (1-9)
                        try:
                            level = int(para.style.name[-1])
                            if 1 <= level <= 6:
                                fontSize = 22 - (level * 2)  # Size decreases with level
                                fontWeight = QFont.Bold
                            else:
                                fontSize = 12
                                fontWeight = QFont.Normal
                        except:
                            fontSize = 16
                            fontWeight = QFont.Bold

                        headingFont = QFont("Times New Roman", fontSize, fontWeight)
                        paraLabel.setFont(headingFont)

                        # Add some spacing before headings
                        documentLayout.addSpacing(10)
                    else:
                        # Regular paragraph
                        paraFont = QFont("Times New Roman", 12)
                        paraLabel.setFont(paraFont)

                    # Set paragraph text with basic formatting
                    formatted_text = ""
                    for run in para.runs:
                        text = run.text

                        # Apply formatting
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                        if run.underline:
                            text = f"<u>{text}</u>"

                        formatted_text += text

                    paraLabel.setText(formatted_text)
                    paraLabel.setTextFormat(Qt.RichText)
                    paraLabel.setWordWrap(True)
                    documentLayout.addWidget(paraLabel)

                # Process tables
                for table in doc.tables:
                    # Create a frame for the table
                    tableFrame = QFrame()
                    tableFrame.setFrameShape(QFrame.Box)
                    tableFrame.setStyleSheet("border: 1px solid #CCCCCC; background-color: #F9F9F9;")

                    # Create grid layout for table
                    tableLayout = QGridLayout(tableFrame)
                    tableLayout.setSpacing(2)

                    # Add table cells
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            # Create cell content
                            cellLabel = QLabel()

                            # Format cell content
                            cell_text = ""
                            for para in cell.paragraphs:
                                if cell_text:
                                    cell_text += "<br>"

                                for run in para.runs:
                                    text = run.text
                                    if run.bold:
                                        text = f"<strong>{text}</strong>"
                                    if run.italic:
                                        text = f"<em>{text}</em>"
                                    if run.underline:
                                        text = f"<u>{text}</u>"

                                    cell_text += text

                            cellLabel.setText(cell_text)
                            cellLabel.setTextFormat(Qt.RichText)
                            cellLabel.setWordWrap(True)
                            cellLabel.setAlignment(Qt.AlignTop | Qt.AlignLeft)

                            # Style header row differently
                            if i == 0:
                                cellLabel.setStyleSheet("font-weight: bold; background-color: #E6E6E6; padding: 5px;")
                            else:
                                cellLabel.setStyleSheet("padding: 5px;")

                            tableLayout.addWidget(cellLabel, i, j)

                    # Add the table to the document layout
                    documentLayout.addWidget(tableFrame)

                # Add images from the document
                self.add_document_images(doc, documentLayout)

                # Add spacing at the end
                documentLayout.addSpacing(20)

        except Exception as e:
            print(f"Error displaying document: {str(e)}")
            errorLabel = QLabel(f"Error displaying document: {str(e)}")
            errorLabel.setFont(QFont("Times New Roman", 12))
            errorLabel.setAlignment(Qt.AlignCenter)
            errorLabel.setWordWrap(True)
            documentLayout.addWidget(errorLabel)

        # Add stretch to push everything to the top
        documentLayout.addStretch()

        # Set the content widget to the scroll area
        scrollArea.setWidget(contentWidget)

        # Adding the widgets to the mainLayout
        mainLayout.addWidget(navBarWidget)
        mainLayout.addWidget(fileSpecWidget)
        mainLayout.addWidget(scrollArea)

        # Setting the main layout to the window
        self.setLayout(mainLayout)

    def add_document_images(self, doc, layout):
        """Extract and display images from the document"""
        try:
            image_index = 0

            # Process all relationships to find images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get the image data
                        image_data = rel.target_part.blob

                        # Create an image from the binary data
                        image = Image.open(io.BytesIO(image_data))

                        # Convert PIL Image to QPixmap
                        img_qt = ImageQt.ImageQt(image)
                        pixmap = QPixmap.fromImage(img_qt)

                        # Create an image label
                        imgLabel = QLabel()
                        imgLabel.setPixmap(pixmap)

                        # Scale the image to fit within the document width while maintaining aspect ratio
                        max_width = 600  # Maximum width in pixels
                        if pixmap.width() > max_width:
                            imgLabel.setPixmap(pixmap.scaledToWidth(max_width, Qt.SmoothTransformation))

                        # Center the image
                        imgLabel.setAlignment(Qt.AlignCenter)

                        # Add to layout
                        layout.addWidget(imgLabel)

                        # Add caption
                        captionLabel = QLabel(f"Figure {image_index + 1}")
                        captionLabel.setAlignment(Qt.AlignCenter)
                        captionLabel.setFont(QFont("Times New Roman", 10, QFont.Bold))
                        layout.addWidget(captionLabel)

                        image_index += 1
                    except Exception as e:
                        print(f"Error processing image {image_index}: {str(e)}")
        except Exception as e:
            print(f"Error adding images: {str(e)}")

    # this function navigates back to the main page
    def navigateHome(self):
        self.stack.setCurrentIndex(1)


# this contains the GUI for opening up a transcript
class TranscriptPage(DocumentPage):
    def __init__(self, stack, fileName=None):
        # Call the parent class constructor
        super().__init__(stack, fileName)

        # Override the fileSystem with a new one pointing to the transcript directory
        self.fileSystem = FileSystem()
        self.fileSystem.fileDirectory = "MeetingTranscripts"

        # Update the file size display to use the transcript size method
        try:
            # Update the file size label
            for i in range(self.layout().count()):
                widget = self.layout().itemAt(i).widget()
                if isinstance(widget, QWidget) and widget.styleSheet().find("background-color: #D1D6D7") != -1:
                    for j in range(widget.layout().count()):
                        childWidget = widget.layout().itemAt(j).widget()
                        if isinstance(childWidget, QLabel) and childWidget.text().startswith("Size:"):
                            fileSize = self.fileSystem.getTranscriptSize(fileName)
                            childWidget.setText(f"Size: {fileSize}")
                            break
                    break
        except Exception as e:
            print(f"Error updating transcript file size: {str(e)}")

        # Reload the document with the correct path
        self.reloadDocument()

    def reloadDocument(self):
        """Reload the document with the correct file path"""
        # Clear the existing content
        contentWidget = None
        scrollArea = None

        # Find the scroll area and content widget
        for i in range(self.layout().count()):
            widget = self.layout().itemAt(i).widget()
            if isinstance(widget, QScrollArea):
                scrollArea = widget
                contentWidget = widget.widget()
                break

        if not contentWidget or not scrollArea:
            return

        # Clear the existing layout
        while contentWidget.layout().count():
            item = contentWidget.layout().takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

        # Reload the document
        documentLayout = contentWidget.layout()

        try:
            # Get the full path to the document
            fullPath = os.path.join(self.fileSystem.fileDirectory, self.fileName)

            # Check if file exists
            if not os.path.exists(fullPath):
                errorLabel = QLabel(f"Error: File not found at {fullPath}")
                errorLabel.setAlignment(Qt.AlignCenter)
                documentLayout.addWidget(errorLabel)
            else:
                # Load the document using python-docx
                doc = Document(fullPath)

                # Add a horizontal line at the top of the document content
                line = QFrame()
                line.setFrameShape(QFrame.HLine)
                line.setFrameShadow(QFrame.Sunken)
                line.setStyleSheet("border: 1px solid #cccccc;")
                documentLayout.addWidget(line)

                # Add some spacing after the line
                documentLayout.addSpacing(10)

                # Process paragraphs
                for para in doc.paragraphs:
                    if not para.text.strip():  # Skip empty paragraphs
                        continue

                    # Create paragraph label
                    paraLabel = QLabel()

                    # Apply paragraph style
                    if para.style.name.startswith('Heading'):
                        # Get heading level (1-9)
                        try:
                            level = int(para.style.name[-1])
                            if 1 <= level <= 6:
                                fontSize = 22 - (level * 2)  # Size decreases with level
                                fontWeight = QFont.Bold
                            else:
                                fontSize = 12
                                fontWeight = QFont.Normal
                        except:
                            fontSize = 16
                            fontWeight = QFont.Bold

                        headingFont = QFont("Times New Roman", fontSize, fontWeight)
                        paraLabel.setFont(headingFont)

                        # Add some spacing before headings
                        documentLayout.addSpacing(10)
                    else:
                        # Regular paragraph
                        paraFont = QFont("Times New Roman", 12)
                        paraLabel.setFont(paraFont)

                    # Set paragraph text with basic formatting
                    formatted_text = ""
                    for run in para.runs:
                        text = run.text

                        # Apply formatting
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                        if run.underline:
                            text = f"<u>{text}</u>"

                        formatted_text += text

                    paraLabel.setText(formatted_text)
                    paraLabel.setTextFormat(Qt.RichText)
                    paraLabel.setWordWrap(True)
                    documentLayout.addWidget(paraLabel)

                # Process tables
                for table in doc.tables:
                    # Create a frame for the table
                    tableFrame = QFrame()
                    tableFrame.setFrameShape(QFrame.Box)
                    tableFrame.setStyleSheet("border: 1px solid #CCCCCC; background-color: #F9F9F9;")

                    # Create grid layout for table
                    tableLayout = QGridLayout(tableFrame)
                    tableLayout.setSpacing(2)

                    # Add table cells
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            # Create cell content
                            cellLabel = QLabel()

                            # Format cell content
                            cell_text = ""
                            for para in cell.paragraphs:
                                if cell_text:
                                    cell_text += "<br>"

                                for run in para.runs:
                                    text = run.text
                                    if run.bold:
                                        text = f"<strong>{text}</strong>"
                                    if run.italic:
                                        text = f"<em>{text}</em>"
                                    if run.underline:
                                        text = f"<u>{text}</u>"

                                    cell_text += text

                            cellLabel.setText(cell_text)
                            cellLabel.setTextFormat(Qt.RichText)
                            cellLabel.setWordWrap(True)
                            cellLabel.setAlignment(Qt.AlignTop | Qt.AlignLeft)

                            # Style header row differently
                            if i == 0:
                                cellLabel.setStyleSheet("font-weight: bold; background-color: #E6E6E6; padding: 5px;")
                            else:
                                cellLabel.setStyleSheet("padding: 5px;")

                            tableLayout.addWidget(cellLabel, i, j)

                    # Add the table to the document layout
                    documentLayout.addWidget(tableFrame)

                # Add images from the document
                self.add_document_images(doc, documentLayout)

                # Add spacing at the end
                documentLayout.addSpacing(20)

        except Exception as e:
            print(f"Error displaying transcript: {str(e)}")
            errorLabel = QLabel(f"Error displaying transcript: {str(e)}")
            errorLabel.setFont(QFont("Times New Roman", 12))
            errorLabel.setAlignment(Qt.AlignCenter)
            errorLabel.setWordWrap(True)
            documentLayout.addWidget(errorLabel)

        # Add stretch to push everything to the top
        documentLayout.addStretch()


# this class manages the stacked pages
class MainWindow(QMainWindow):

    def __init__(self):
        super().__init__()

        # The QStackedLayout will store all the different QWidgets/pages
        self.stack = QStackedWidget()
        
        # storing all the pages to the stack
        self.loginPage = LoginWindow(self.stack)
        self.mainPage = MainPage(self.stack)
        self.summarizationPage = SummarizationPage(self.stack)
        self.recordAudioPage = RecordAudioPage(self.stack)

        # The way the pages are added to the stack determine thier index
        self.stack.addWidget(self.loginPage) # 0
        self.stack.addWidget(self.mainPage) # 1
        self.stack.addWidget(self.summarizationPage) # 2
        self.stack.addWidget(self.recordAudioPage) # 2


        # Set the stacked widget as the central widget of QMainWindow
        self.setCentralWidget(self.stack)

        # setting up the default size of the browser
        self.setMinimumSize(900, 600)

# config setup for Qapplication
app = QApplication(sys.argv)

window = MainWindow()
window.show()

sys.exit(app.exec_())
