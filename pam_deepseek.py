import re
import os
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import docx2txt
from docx import Document
import random
import numpy as np
import time
from tqdm import tqdm


class ProgressTracker:
    """Class to track and display progress of transcript processing"""

    def __init__(self, total_steps=5):
        self.total_steps = total_steps
        self.current_step = 0
        self.step_progress = 0
        self.start_time = time.time()
        self.step_start_time = time.time()

    def update_step(self, step_name, step_progress=0):
        """Update to a new processing step"""
        self.current_step += 1
        self.step_progress = step_progress
        self.step_start_time = time.time()
        elapsed = time.time() - self.start_time
        overall_progress = min(100, (self.current_step - 1 + step_progress / 100) / self.total_steps * 100)

        print(
            f"[{overall_progress:.1f}%] Step {self.current_step}/{self.total_steps}: {step_name} | Time elapsed: {elapsed:.1f}s")

    def update_progress(self, progress, items_processed=None, total_items=None):
        """Update progress within the current step"""
        self.step_progress = progress
        elapsed = time.time() - self.start_time
        step_elapsed = time.time() - self.step_start_time
        overall_progress = min(100, (self.current_step - 1 + progress / 100) / self.total_steps * 100)

        if items_processed is not None and total_items is not None:
            print(
                f"[{overall_progress:.1f}%] Step {self.current_step}/{self.total_steps}: {items_processed}/{total_items} items processed | Time elapsed: {elapsed:.1f}s")
        else:
            print(f"[{overall_progress:.1f}%] Current step: {progress:.1f}% complete | Time elapsed: {elapsed:.1f}s")

    def complete(self):
        """Mark processing as complete"""
        total_time = time.time() - self.start_time
        print(f"[100.0%] Processing complete! Total time: {total_time:.1f}s")


def chunk_dialogue(dialogue, chunk_size=20, overlap=5):
    """
    Split dialogue into smaller chunks with overlap to maintain context

    Args:
        dialogue: List of dialogue entries
        chunk_size: Number of entries per chunk
        overlap: Number of entries to overlap between chunks

    Returns:
        List of dialogue chunks
    """
    if not dialogue:
        return []

    if len(dialogue) <= chunk_size:
        return [dialogue]

    chunks = []
    i = 0

    while i < len(dialogue):
        end_idx = min(i + chunk_size, len(dialogue))
        chunks.append(dialogue[i:end_idx])
        i += chunk_size - overlap

    return chunks


def estimate_token_count(text, tokenizer):
    """Estimate token count for a given text"""
    return len(tokenizer.encode(text))


def filter_off_topic_conversation(dialogue, tokenizer, model, progress_tracker=None):
    """Filter out off-topic conversations from the dialogue using chunking for large transcripts"""
    # If dialogue is empty, return as is
    if not dialogue:
        return dialogue

    if progress_tracker:
        progress_tracker.update_step("Identifying meeting topics")

    # Extract the first few entries to identify meeting topics
    initial_entries = dialogue[:min(10, len(dialogue))]
    initial_content = "\n".join([f"{entry['speaker']}: {entry['content']}" for entry in initial_entries])

    # Create prompt to identify main topics
    topic_prompt = [
        {"role": "system", "content": "You are an assistant that helps identify the main topics of a meeting."},
        {"role": "user", "content": f"""Based on the beginning of this meeting transcript, identify the 3-5 main topics that are likely to be discussed in this meeting. Be specific and concise.

Transcript beginning:
{initial_content}

Respond with only the topics, one per line, without numbering or prefixes."""}
    ]

    try:
        # Generate topics
        inputs = tokenizer.apply_chat_template(
            topic_prompt,
            add_generation_prompt=True,
            return_tensors="pt"
        ).to(model.device)

        attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

        outputs = model.generate(
            inputs,
            attention_mask=attention_mask,
            max_new_tokens=200,
            temperature=0.7,
            do_sample=True,
            pad_token_id=tokenizer.eos_token_id
        )

        topics = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True).strip()
        topic_list = [topic.strip() for topic in topics.split('\n') if topic.strip()]

        # If no topics were identified, return original dialogue
        if not topic_list:
            return dialogue

    except Exception as e:
        print(f"Error identifying topics: {e}")
        return dialogue  # On error, return the original dialogue

    if progress_tracker:
        progress_tracker.update_step("Filtering dialogue entries")

    # Process dialogue in chunks to avoid token limit issues
    filtered_dialogue = []
    chunks = chunk_dialogue(dialogue, chunk_size=15, overlap=3)

    for chunk_idx, chunk in enumerate(chunks):
        # Update progress
        if progress_tracker:
            chunk_progress = (chunk_idx / len(chunks)) * 100
            progress_tracker.update_progress(chunk_progress, chunk_idx + 1, len(chunks))

        try:
            chunk_entries_text = []
            for entry in chunk:
                chunk_entries_text.append(f"{entry['speaker']}: {entry['content']}")

            chunk_text = "\n".join(chunk_entries_text)

            # Estimate token count to avoid exceeding limits
            prompt_text = f"""The meeting is focused on these main topics:
{', '.join(topic_list)}

For each of the following dialogue entries, determine if it is ON-TOPIC (relevant to the meeting topics) or OFF-TOPIC (irrelevant, social chat, etc.)

{chunk_text}

For each entry, respond ONLY with "ON-TOPIC" or "OFF-TOPIC", one answer per line, nothing else."""

            estimated_tokens = estimate_token_count(prompt_text, tokenizer)
            if estimated_tokens > 3500:  # Safe limit below the 4096 max
                # If too long, further break down the chunk
                sub_chunks = chunk_dialogue(chunk, chunk_size=5, overlap=1)
                chunk_results = []

                for sub_chunk in sub_chunks:
                    sub_entries_text = []
                    for entry in sub_chunk:
                        sub_entries_text.append(f"{entry['speaker']}: {entry['content']}")

                    sub_text = "\n".join(sub_entries_text)

                    filter_prompt = [
                        {"role": "system",
                         "content": "You are an assistant that helps identify relevant parts of a meeting transcript."},
                        {"role": "user", "content": f"""The meeting is focused on these main topics:
{', '.join(topic_list)}

For each of the following dialogue entries, determine if it is ON-TOPIC (relevant to the meeting topics) or OFF-TOPIC (irrelevant, social chat, etc.)

{sub_text}

For each entry, respond ONLY with "ON-TOPIC" or "OFF-TOPIC", one answer per line, nothing else."""}
                    ]

                    inputs = tokenizer.apply_chat_template(
                        filter_prompt,
                        add_generation_prompt=True,
                        return_tensors="pt"
                    ).to(model.device)

                    attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

                    outputs = model.generate(
                        inputs,
                        attention_mask=attention_mask,
                        max_new_tokens=len(sub_chunk) * 15,
                        temperature=0.2,
                        do_sample=True,
                        pad_token_id=tokenizer.eos_token_id
                    )

                    sub_decisions = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True).strip()
                    chunk_results.extend([d.strip() for d in sub_decisions.split('\n') if d.strip()])

                decisions = chunk_results
            else:
                # Process the entire chunk if it's within token limits
                filter_prompt = [
                    {"role": "system",
                     "content": "You are an assistant that helps identify relevant parts of a meeting transcript."},
                    {"role": "user", "content": prompt_text}
                ]

                inputs = tokenizer.apply_chat_template(
                    filter_prompt,
                    add_generation_prompt=True,
                    return_tensors="pt"
                ).to(model.device)

                attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

                outputs = model.generate(
                    inputs,
                    attention_mask=attention_mask,
                    max_new_tokens=len(chunk) * 15,  # Allow enough tokens for responses
                    temperature=0.2,  # Low temperature for more consistent classification
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id
                )

                decisions = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True).strip()
                decisions = [d.strip() for d in decisions.split('\n') if d.strip()]

            # Ensure decisions align with chunk entries
            if len(decisions) < len(chunk):
                # If we don't have enough decisions, keep all entries to be safe
                filtered_chunk = chunk
            else:
                # Keep only on-topic entries
                filtered_chunk = []
                for i, entry in enumerate(chunk):
                    if i < len(decisions) and "ON-TOPIC" in decisions[i].upper():
                        filtered_chunk.append(entry)
                    # Always keep the first few entries of each chunk to maintain context
                    elif i < min(3, len(chunk)) and chunk_idx == 0:
                        filtered_chunk.append(entry)

            # For chunks after the first, check for overlapping entries to avoid duplicates
            if chunk_idx > 0 and filtered_dialogue:
                # Get IDs or unique identifiers for already processed entries
                existing_ids = {f"{e['speaker']}-{e['timestamp']}": True for e in filtered_dialogue}

                # Only add non-duplicate entries from this chunk
                for entry in filtered_chunk:
                    entry_id = f"{entry['speaker']}-{entry['timestamp']}"
                    if entry_id not in existing_ids:
                        filtered_dialogue.append(entry)
            else:
                # For the first chunk, add all filtered entries
                filtered_dialogue.extend(filtered_chunk)

        except Exception as e:
            print(f"Error processing chunk {chunk_idx + 1}/{len(chunks)}: {e}")
            # On error, include all entries from this chunk to be safe
            filtered_dialogue.extend(chunk)

    # If filtering removed too much, return original dialogue
    if len(filtered_dialogue) < len(dialogue) * 0.3:  # If we removed more than 70%
        return dialogue

    return filtered_dialogue


def generate_fallback_minutes(dialogue, attendees, date_time):
    """Generate basic meeting minutes when the main generation process fails"""
    # Extract basic information from dialogue
    topics = []
    decisions = []
    action_items = []
    next_steps = []

    # Basic extraction of potential topics, decisions, and action items based on keywords
    for entry in dialogue:
        content = entry["content"].lower()

        # Look for topic indicators
        if any(kw in content for kw in ["agenda", "topic", "discuss", "talking about", "focus on"]):
            topics.append(f"- Discussion related to {entry['content'][:50]}...")

        # Look for decision indicators
        if any(kw in content for kw in ["decide", "decision", "agreed", "agreement", "conclusion", "resolved"]):
            decisions.append(f"- {entry['speaker']} mentioned: {entry['content'][:50]}...")

        # Look for action item indicators
        if any(kw in content for kw in ["action", "task", "todo", "to-do", "assign", "responsibility", "will do"]):
            action_items.append(f"- {entry['speaker']} will {entry['content'][:50]}...")

        # Look for next steps indicators
        if any(kw in content for kw in ["next step", "follow up", "following up", "next time", "next meeting"]):
            next_steps.append(f"- {entry['content'][:50]}...")

    # Deduplicate and limit each section
    topics = list(set(topics))[:5]
    decisions = list(set(decisions))[:5]
    action_items = list(set(action_items))[:5]
    next_steps = list(set(next_steps))[:5]

    # Create fallback minutes document
    minutes = f"""# Meeting Minutes

## Date and Time
{date_time}

## Attendees
{', '.join(attendees)}

## Key Discussion Topics
"""

    if topics:
        for topic in topics:
            minutes += f"{topic}\n"
    else:
        minutes += "- No specific topics were identified automatically\n"

    minutes += "\n## Decisions Made\n"

    if decisions:
        for decision in decisions:
            minutes += f"{decision}\n"
    else:
        minutes += "- No specific decisions were recorded in this meeting\n"

    minutes += "\n## Action Items\n"

    if action_items:
        for item in action_items:
            minutes += f"{item}\n"
    else:
        minutes += "- No specific action items were assigned in this meeting\n"

    minutes += "\n## Next Steps\n"

    if next_steps:
        for step in next_steps:
            minutes += f"{step}\n"
    else:
        minutes += "- No specific next steps were outlined in this meeting\n"

    return minutes

def process_transcript(transcript_path, output_dir="MeetingNotes"):
    """Process transcript and generate clean meeting minutes"""
    # Initialize progress tracker
    progress_tracker = ProgressTracker(total_steps=5)
    progress_tracker.update_step("Initializing model and tokenizer")

    # Seed configuration
    seed = 42
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    torch.cuda.manual_seed_all(seed)
    torch.backends.cudnn.deterministic = True
    torch.backends.cudnn.benchmark = False

    # Load model and tokenizer
    model_name = "DeepSeek-R1-Distill-Qwen-1.5B"
    tokenizer = AutoTokenizer.from_pretrained(model_name)

    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token
        tokenizer.pad_token_id = tokenizer.eos_token_id

    model = AutoModelForCausalLM.from_pretrained(
        model_name,
        torch_dtype=torch.float16,
        device_map="auto",
        pad_token_id=tokenizer.eos_token_id
    )

    # Parse transcript
    progress_tracker.update_step("Parsing transcript")
    try:
        transcript_text = docx2txt.process(transcript_path)
        dialogue = parse_dialogue(transcript_text)
        progress_tracker.update_progress(100)  # 100% of parsing step complete
    except Exception as e:
        print(f"Error parsing transcript: {e}")
        return None

    # Extract relevant information from dialogue
    progress_tracker.update_step("Extracting metadata")
    try:
        attendees = extract_attendees(dialogue)
        date_time = extract_date_time(dialogue)
        progress_tracker.update_progress(100)  # 100% of metadata extraction complete
    except Exception as e:
        print(f"Error extracting metadata: {e}")
        attendees = []
        date_time = "Unknown date and time"

    # Filter out off-topic conversations
    try:
        filtered_dialogue = filter_off_topic_conversation(dialogue, tokenizer, model, progress_tracker)
    except Exception as e:
        print(f"Error filtering dialogue: {e}")
        filtered_dialogue = dialogue  # Use original dialogue if filtering fails

    # Generate minutes
    progress_tracker.update_step("Generating minutes")
    try:
        minutes = generate_clean_minutes(filtered_dialogue, tokenizer, model, attendees, date_time, progress_tracker)
    except Exception as e:
        print(f"Error generating minutes: {e}")
        # Create basic minutes as fallback
        minutes = generate_fallback_minutes(filtered_dialogue, attendees, date_time)

    # Save output
    progress_tracker.update_step("Saving document")
    try:
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, os.path.basename(transcript_path).split('.')[0] + "_minutes.docx")
        save_as_docx(minutes, output_path)
        progress_tracker.update_progress(100)  # 100% of saving complete
    except Exception as e:
        print(f"Error saving document: {e}")
        return None

    progress_tracker.complete()
    return output_path

def parse_dialogue(text):
    """Improved dialogue parser with timestamp handling"""
    entries = []
    pattern = r"(\d+:\d+:\d+\.\d+)\s*-->\s*(\d+:\d+:\d+\.\d+)\s*([\r\n]+)(.*?)\s*([\r\n]+)(.*?)(?=\d+:\d+:\d+\.\d+|$)"
    matches = list(re.finditer(pattern, text, re.DOTALL))

    for i, match in enumerate(matches):
        speaker = match.group(4).strip()
        content = match.group(6).strip()
        timestamp = match.group(1).strip()
        if speaker and content:
            entries.append({"speaker": speaker, "content": content, "timestamp": timestamp})

    return entries


def extract_attendees(dialogue):
    """Extract unique attendees from dialogue"""
    attendees = set()
    for entry in dialogue:
        attendees.add(entry["speaker"])
    return sorted(list(attendees))


def extract_date_time(dialogue):
    """Extract date and time information if available"""
    # Look for date/time patterns in the dialogue
    date_pattern = r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(?:st|nd|rd|th)?,\s+\d{4}\b"
    time_pattern = r"\b\d{1,2}:\d{2}\s*(?:AM|PM)\b"

    # Initialize with default
    date_time_info = "No specific date or time mentioned in the transcript"

    # Check all dialogue for dates
    full_text = " ".join([entry["content"] for entry in dialogue])
    date_match = re.search(date_pattern, full_text, re.IGNORECASE)
    time_match = re.search(time_pattern, full_text, re.IGNORECASE)

    if date_match and time_match:
        date_time_info = f"Meeting held on {date_match.group(0)} at {time_match.group(0)}"
    elif date_match:
        date_time_info = f"Meeting held on {date_match.group(0)}"
    elif time_match:
        date_time_info = f"Meeting held at {time_match.group(0)}"

    return date_time_info


def generate_clean_minutes(dialogue, tokenizer, model, attendees, date_time, progress_tracker=None):
    """Generate structured minutes with chunking for large transcripts"""
    if progress_tracker:
        progress_tracker.update_progress(10, "Preparing transcript for processing")

    # For large transcripts, process in chunks
    if len(dialogue) > 30:  # Threshold for chunking
        return generate_chunked_minutes(dialogue, tokenizer, model, attendees, date_time, progress_tracker)

    # For smaller transcripts, use the original approach
    transcript_chunk = "\n".join([f"{entry['speaker']}: {entry['content']}" for entry in dialogue])

    # Create a more controlled prompt with clear instructions
    system_prompt = ("You are an expert meeting minutes writer. Your task is to create professional meeting minutes "
                     "that accurately summarize the key points from the transcript. Extract only factual information. "
                     "Do not include any commentary, thinking process, notes to yourself, or concluding statements. "
                     "Never invent information that isn't explicitly in the transcript.")

    user_prompt = f"""Based on the meeting transcript below, create concise and professional meeting minutes following this exact structure:

# Meeting Minutes

## Date and Time
{date_time}

## Attendees
{', '.join(attendees)}

## Key Discussion Topics
- [Topic 1]
- [Topic 2]
- [Topic 3]
(Use bullet points/dashes, NOT numbered items)

## Decisions Made
- [Decision 1]
- [Decision 2]
(Use bullet points/dashes, NOT numbered items)

## Action Items
- [Task 1]
- [Task 2]
(Use bullet points/dashes, NOT numbered items)

## Next Steps
- [Step 1]
- [Step 2]
(Use bullet points/dashes, NOT numbered items)

Transcript:
{transcript_chunk}

IMPORTANT: 
1. Do NOT use numbered lists anywhere in the document. Always use bullet points with dashes (-).
2. Do NOT add any text, summary, or conclusion after the "Next Steps" section.
3. Do NOT include any thinking process, notes to yourself, or placeholder text.
4. Each section should contain only extracted information from the transcript.
5. If information for a section is not available, provide a brief factual statement like "No specific decisions were recorded in this meeting."
6. End the document immediately after the last item in the "Next Steps" section.
7. CRUCIAL: Do NOT invent or hallucinate any information that isn't explicitly mentioned in the transcript, especially dates, times, or numerical data. If information is not provided, use exactly what's in the "Date and Time" section I've provided.
"""

    if progress_tracker:
        progress_tracker.update_progress(20, "Creating prompt")

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    # Check estimated token count
    full_prompt = system_prompt + user_prompt
    estimated_tokens = estimate_token_count(full_prompt, tokenizer)

    if estimated_tokens > 3500:  # If too large
        # Truncate transcript but keep structure
        max_tokens = 2500  # Target max tokens for the prompt
        system_and_structure_tokens = estimate_token_count(system_prompt + user_prompt.split("Transcript:")[0],
                                                           tokenizer)
        available_tokens = max_tokens - system_and_structure_tokens

        # Create a shorter transcript
        shortened_transcript = ""
        current_tokens = 0

        for entry in dialogue:
            entry_text = f"{entry['speaker']}: {entry['content']}\n"
            entry_tokens = estimate_token_count(entry_text, tokenizer)

            if current_tokens + entry_tokens <= available_tokens:
                shortened_transcript += entry_text
                current_tokens += entry_tokens
            else:
                break

        # Update user prompt with shortened transcript
        user_prompt = user_prompt.split("Transcript:")[
                          0] + f"Transcript:\n{shortened_transcript}\n\nNote: This is a truncated portion of the meeting transcript."
        messages[1]["content"] = user_prompt

    if progress_tracker:
        progress_tracker.update_progress(30, "Tokenizing input")

    try:
        # Tokenize with attention mask
        inputs = tokenizer.apply_chat_template(
            messages,
            add_generation_prompt=True,
            return_tensors="pt"
        ).to(model.device)

        attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

        if progress_tracker:
            progress_tracker.update_progress(50, "Generating content")

        # Generate with proper parameters
        outputs = model.generate(
            inputs,
            attention_mask=attention_mask,
            max_new_tokens=1000,
            repetition_penalty=1.1,
            temperature=0.6,  # Slightly lower temperature for more factual output
            top_p=0.92,
            do_sample=True,
            pad_token_id=tokenizer.eos_token_id
        )

        # Decode and clean
        if progress_tracker:
            progress_tracker.update_progress(80, "Post-processing output")

        full_output = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True)
        processed_output = post_process_output(full_output)

    except Exception as e:
        print(f"Error generating minutes: {e}")
        return generate_fallback_minutes(dialogue, attendees, date_time)

    if progress_tracker:
        progress_tracker.update_progress(100, "Minutes generation complete")

    return processed_output

def generate_chunked_minutes(dialogue, tokenizer, model, attendees, date_time, progress_tracker=None):
    """Process a large transcript in chunks and combine the results"""
    if progress_tracker:
        progress_tracker.update_progress(15, "Breaking transcript into manageable chunks")

    # Step 1: First generate basic summary from the beginning and end of the transcript
    intro_section = dialogue[:min(15, len(dialogue))]
    outro_section = dialogue[max(0, len(dialogue) - 15):]

    # Combine intro and outro for overview
    overview_dialogue = intro_section + outro_section
    overview_text = "\n".join([f"{entry['speaker']}: {entry['content']}" for entry in overview_dialogue])

    # Step 2: Generate overview minutes to establish structure
    overview_system_prompt = (
        "You are an expert meeting minutes writer. Analyze the beginning and end of this meeting transcript "
        "to identify the main purpose, key participants, and overall structure.")

    overview_user_prompt = f"""Based on the transcript excerpts below (from the beginning and end of the meeting), identify:
1. The main purpose of the meeting
2. The key discussion areas (3-5 topics)
3. Any mentioned decisions or action items

Transcript excerpts:
{overview_text}

Keep your response brief and focused on what's explicitly mentioned. Do not invent details.
"""

    try:
        overview_messages = [
            {"role": "system", "content": overview_system_prompt},
            {"role": "user", "content": overview_user_prompt}
        ]

        inputs = tokenizer.apply_chat_template(
            overview_messages,
            add_generation_prompt=True,
            return_tensors="pt"
        ).to(model.device)

        attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

        outputs = model.generate(
            inputs,
            attention_mask=attention_mask,
            max_new_tokens=500,
            temperature=0.5,
            do_sample=True,
            pad_token_id=tokenizer.eos_token_id
        )

        overview = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True)
    except Exception as e:
        print(f"Error generating overview: {e}")
        overview = "Unable to generate meeting overview."

    # Step 3: Break dialogue into chunks and process each chunk
    chunks = chunk_dialogue(dialogue, chunk_size=20, overlap=3)

    # Containers for section content
    topics = []
    decisions = []
    action_items = []
    next_steps = []

    if progress_tracker:
        progress_tracker.update_progress(30, "Processing transcript chunks")

    # Process each chunk to extract relevant information
    for chunk_idx, chunk in enumerate(chunks):
        if progress_tracker:
            chunk_progress = 30 + (50 * chunk_idx / len(chunks))
            progress_tracker.update_progress(chunk_progress, chunk_idx + 1, len(chunks))

        chunk_text = "\n".join([f"{entry['speaker']}: {entry['content']}" for entry in chunk])

        # Create a targeted prompt for extracting specific information from this chunk
        chunk_system_prompt = "You are an assistant that extracts specific information from meeting transcript segments."

        chunk_user_prompt = f"""From the following meeting transcript chunk, extract ONLY:
1. Key discussion topics
2. Any decisions made
3. Action items assigned
4. Next steps mentioned

For each category, list only the items explicitly mentioned in this chunk.
Format your response with these exact headings:

TOPICS:
- [topic 1]
- [topic 2]

DECISIONS:
- [decision 1]
- [decision 2]

ACTION ITEMS:
- [action 1]
- [action 2]

NEXT STEPS:
- [step 1]
- [step 2]

If nothing is mentioned for a category, write "None mentioned in this segment." under the heading.

Transcript chunk:
{chunk_text}
"""

        try:
            chunk_messages = [
                {"role": "system", "content": chunk_system_prompt},
                {"role": "user", "content": chunk_user_prompt}
            ]

            inputs = tokenizer.apply_chat_template(
                chunk_messages,
                add_generation_prompt=True,
                return_tensors="pt"
            ).to(model.device)

            attention_mask = inputs.ne(tokenizer.pad_token_id).float().to(model.device)

            outputs = model.generate(
                inputs,
                attention_mask=attention_mask,
                max_new_tokens=400,
                temperature=0.4,
                do_sample=True,
                pad_token_id=tokenizer.eos_token_id
            )

            chunk_results = tokenizer.decode(outputs[0][inputs.shape[1]:], skip_special_tokens=True)

            # Extract information from chunk results
            topic_section = re.search(r"TOPICS:(.*?)(?=DECISIONS:|$)", chunk_results, re.DOTALL)
            decision_section = re.search(r"DECISIONS:(.*?)(?=ACTION ITEMS:|$)", chunk_results, re.DOTALL)
            action_section = re.search(r"ACTION ITEMS:(.*?)(?=NEXT STEPS:|$)", chunk_results, re.DOTALL)
            next_steps_section = re.search(r"NEXT STEPS:(.*?)$", chunk_results, re.DOTALL)

            # Add extracted items to their respective containers
            if topic_section:
                for line in topic_section.group(1).strip().split('\n'):
                    line = line.strip()
                    if line and line.startswith('-') and "None mentioned" not in line:
                        topics.append(line)

            if decision_section:
                for line in decision_section.group(1).strip().split('\n'):
                    line = line.strip()
                    if line and line.startswith('-') and "None mentioned" not in line:
                        decisions.append(line)

            if action_section:
                for line in action_section.group(1).strip().split('\n'):
                    line = line.strip()
                    if line and line.startswith('-') and "None mentioned" not in line:
                        action_items.append(line)

            if next_steps_section:
                for line in next_steps_section.group(1).strip().split('\n'):
                    line = line.strip()
                    if line and line.startswith('-') and "None mentioned" not in line:
                        next_steps.append(line)

        except Exception as e:
            print(f"Error processing chunk {chunk_idx + 1}/{len(chunks)}: {e}")

    # Step 4: Remove duplicates while preserving order
    def deduplicate(items):
        seen = set()
        result = []
        for item in items:
            normalized = ' '.join(item.lower().split())
            if normalized not in seen and normalized:
                seen.add(normalized)
                result.append(item)
        return result

    topics = deduplicate(topics)
    decisions = deduplicate(decisions)
    action_items = deduplicate(action_items)
    next_steps = deduplicate(next_steps)

    if progress_tracker:
        progress_tracker.update_progress(85, "Compiling final minutes")

    # Step 5: Compile the final minutes document
    minutes = f"""# Meeting Minutes

## Date and Time
{date_time}

## Attendees
{', '.join(attendees)}

## Key Discussion Topics
"""

    if topics:
        for topic in topics[:5]:  # Limit to 5 most important topics
            minutes += f"{topic}\n"
    else:
        minutes += "- No specific topics were clearly identified in the transcript\n"

    minutes += "\n## Decisions Made\n"

    if decisions:
        for decision in decisions:
            minutes += f"{decision}\n"
    else:
        minutes += "- No specific decisions were recorded in this meeting\n"

    minutes += "\n## Action Items\n"

    if action_items:
        for item in action_items:
            minutes += f"{item}\n"
    else:
        minutes += "- No specific action items were assigned in this meeting\n"

    minutes += "\n## Next Steps\n"

    if next_steps:
        for step in next_steps:
            minutes += f"{step}\n"
    else:
        minutes += "- No specific next steps were outlined in this meeting\n"

    return post_process_output(minutes)

def post_process_output(text):
    """Clean and structure the generated output"""
    # Remove any residual prompt fragments or thinking notes
    text = re.sub(r"(Transcript:|IMPORTANT:.*?|</?think>|Here\'s|I\'ll|Let me)", "", text,
                  flags=re.DOTALL | re.IGNORECASE)

    # Remove markdown formatting artifacts that shouldn't be there
    text = re.sub(r"[-\*]{2,}|\\-\\-\\-", "", text)

    # If the output starts with anything other than the main heading, add it
    if not text.strip().startswith("# Meeting Minutes"):
        text = "# Meeting Minutes\n\n" + text

    # Ensure each section exists and appears only once
    required_sections = [
        "## Date and Time",
        "## Attendees",
        "## Key Discussion Topics",
        "## Decisions Made",
        "## Action Items",
        "## Next Steps"
    ]

    # Process one section at a time to ensure correct ordering
    processed_text = "# Meeting Minutes\n\n"

    for section in required_sections:
        pattern = f"{re.escape(section)}(.*?)(?=## |$)"
        matches = re.search(pattern, text, re.DOTALL)

        if matches:
            section_content = matches.group(1).strip()

            # Check for hallucinated dates if this is the Date and Time section
            if section == "## Date and Time" and section_content:
                # Check if the content contains a date that wasn't in the original date_time string
                date_pattern = r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(?:st|nd|rd|th)?,\s+\d{4}\b"
                year_pattern = r"\b20\d{2}\b"

                # If there's a specific date pattern or year that wasn't in the original extraction
                if (re.search(date_pattern, section_content, re.IGNORECASE) or
                        re.search(year_pattern, section_content)):
                    # Replace with a safer message
                    section_content = "No specific date or time mentioned in the transcript"

            if section_content:
                # Convert any numbered lists to bullet points
                section_content = re.sub(r"^\s*\d+\.\s+", "- ", section_content, flags=re.MULTILINE)
                processed_text += f"{section}\n{section_content}\n\n"
            else:
                processed_text += f"{section}\n- Information not available in the transcript\n\n"
        else:
            processed_text += f"{section}\n- Information not available in the transcript\n\n"

    # Clean formatting
    processed_text = re.sub(r"\n{3,}", "\n\n", processed_text)
    processed_text = re.sub(r"(?<!\n)\n(?=#)", "\n\n", processed_text)

    # Convert various bullet point types to consistent format
    processed_text = re.sub(r"^\s*[•\*]\s*", "- ", processed_text, flags=re.MULTILINE)

    # Remove any content after the "Next Steps" section
    next_steps_pattern = r"(## Next Steps.*?)(?:\n\n|$)(.+)?$"
    matches = re.search(next_steps_pattern, processed_text, re.DOTALL)
    if matches and matches.group(2):
        # If there's content after Next Steps, remove it
        processed_text = re.sub(next_steps_pattern, r"\1\n", processed_text, flags=re.DOTALL)

    return processed_text.strip()


def save_as_docx(minutes, output_path):
    """Save formatted document with proper headings"""
    doc = Document()

    lines = minutes.split('\n')
    i = 0
    next_steps_found = False
    next_steps_content_found = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Check if we've reached Next Steps section
        if line == "## Next Steps":
            next_steps_found = True
            # Add the heading
            doc.add_heading(line[3:], level=1)
            i += 1
            continue

        # Skip any content that appears after Next Steps section is complete
        if next_steps_found and next_steps_content_found and not line.startswith("- "):
            i += 1
            continue

        if line.startswith('# '):
            # Main title
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            # Section heading
            doc.add_heading(line[3:], level=1)
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(line[2:])
            # If we're in Next Steps and find bullet content, mark it
            if next_steps_found:
                next_steps_content_found = True
        elif re.match(r'\d+\.\s+', line):
            # Convert numbered list to bullet point
            p = doc.add_paragraph()
            p.style = 'List Bullet'  # Change to bullet, not number
            p.add_run(re.sub(r'\d+\.\s+', '', line))
            # If we're in Next Steps and find bullet content, mark it
            if next_steps_found:
                next_steps_content_found = True
        else:
            # Regular paragraph (but not after Next Steps)
            if not (next_steps_found and next_steps_content_found):
                doc.add_paragraph(line)

        i += 1

    doc.save(output_path)


def estimate_completion_time(transcript_size, steps_completed=0, total_steps=5):
    """Estimate remaining processing time based on transcript size"""
    # These values would need to be calibrated based on actual performance data
    # Currently providing rough estimates
    base_time = 10  # Base processing time in seconds
    per_kb_time = 0.5  # Additional seconds per KB of transcript

    estimated_total = base_time + (transcript_size / 1024) * per_kb_time
    estimated_remaining = estimated_total * (1 - steps_completed / total_steps)

    return estimated_remaining


if __name__ == "__main__":
    transcript_path = "transcript01.docx"
    output_path = process_transcript(transcript_path)
    print(f"Professional meeting minutes generated at: {output_path}")
