import os
import time
import docx2txt
from docx import Document
import re
from datetime import datetime, timedelta

# Clean up transcript, timestamp/speaker = 0 to remove
def cleanTranscript (transcript, timestamp, speaker):
    # split each line of text into list
    splitTranscript = transcript.splitlines(keepends=True)

    # check modes and perform corresponding list item deletions
    if (timestamp == 0) and (speaker == 0):
        del splitTranscript[0::4]
        del splitTranscript[0::3]
    elif (timestamp == 0) and (speaker == 1):
        del splitTranscript[0::4]
    elif (timestamp == 1) and (speaker == 0):
        del splitTranscript[1::4]

    # return text to string format and return it
    text = ""
    for line in splitTranscript:
        if line != '\n':
            text += line
    return text

# Obtain the duration of a meeting
def meetingDuration (filepath):
    transcript = docx2txt.process(filepath)
    subString = ''
    subString2 = ''
    for c in reversed(transcript):
        if c == '>':
            break
        subString += c
    for c in reversed(subString):
        subString2 += c
    subString = subString2[1:]
    if subString[1] == ":":
        durationHours = int(subString[0])
        durationMinutes = int(subString[2:4])
    else:
        durationHours = int(subString[0:2])
        durationMinutes = int(subString[3:5])
    return durationHours, durationMinutes

# Calculate meeting start time using raw file date minus meeting duration
def calculateStart (filepath):
    rawDate = time.ctime(os.path.getctime(filepath))
    duration = meetingDuration(filepath)
    if rawDate[9].isspace():
        rawHours = int(rawDate[10:12])
        rawMinutes = int(rawDate[13:15])
    else:
        rawHours = int(rawDate[11:13])
        rawMinutes = int(rawDate[14:16])
    minutes = rawMinutes - duration[1]
    if minutes < 0:
        minutes += 60
        hours = rawHours - duration[0] - 1
    else:
        hours = rawHours - duration[0]
    return hours, minutes

# Obtain formal form of creation date of transcript file and meeting start time
def extractFormalDate (filepath):
    rawDate = time.ctime(os.path.getctime(filepath))
    startTime = calculateStart(filepath)
    if rawDate[9].isspace():
        date = Date(rawDate[:3], rawDate[4:7], rawDate[8], startTime[0], startTime[1], "", rawDate[19:23])
    else:
        date = Date(rawDate[:3], rawDate[4:7], rawDate[8:10], startTime[0], startTime[1], "", rawDate[20:24])
    return date

# Obtain file date for use in file naming convention
def extractFilenameDate (filepath):
    rawDate = time.ctime(os.path.getctime(filepath))
    date = rawDate[-4:] + "-"
    ampm = ""
    month = rawDate[4:7]
    match month:
        case "Jan":
            date += "01"
        case "Feb":
            date += "02"
        case "Mar":
            date += "03"
        case "Apr":
            date += "04"
        case "May":
            date += "05"
        case "Jun":
            date += "06"
        case "Jul":
            date += "07"
        case "Aug":
            date += "08"
        case "Sep":
            date += "09"
        case "Oct":
            date += "10"
        case "Nov":
            date += "11"
        case "Dec":
            date += "12"
    date += "-" + rawDate[8:10] + "_"
    if rawDate[12] == ":":
        date += rawDate[11] + "-"
        date += rawDate[13:15]
        ampm = "AM"
    elif rawDate[13] == ":":
        if rawDate[11:13] == "24":
            date += "12"
            ampm = "AM"
        elif int(rawDate[11:13]) > 12:
            date += str(int(rawDate[11:13]) - 12)
            ampm = "PM"
        elif rawDate[11:13] == "12":
            date += "12"
            ampm = "PM"
        else:
            date += str(int(rawDate[11:13]))
            ampm = "AM"
        date += "\u2236" + rawDate[14:16]
    date += ampm
    return date

def convertZoom(file_path):
    # Regular expression to parse [Speaker] timestamp format
    pattern = re.compile(r"\[(.*?)\] (\d{2}:\d{2}:\d{2})\n(.*)")

    # Read the transcript file
    with open(file_path, "r", encoding="utf-8") as file:
        lines = file.read().strip().split("\n\n")  # Split each speaker block

    # Process transcript data
    transformed_data = []
    previous_speaker = None
    previous_timestamp = None
    buffer_text = ""

    for block in lines:
        match = pattern.match(block)
        if match:
            speaker, timestamp, text = match.groups()
            current_time = datetime.strptime(timestamp, "%H:%M:%S")

            # If it's the same speaker as the previous, append the text
            if speaker == previous_speaker:
                buffer_text += f" {text}"
            else:
                if previous_speaker:
                    transformed_data.append((previous_speaker, previous_timestamp, current_time, buffer_text.strip()))

                # Reset buffer for the new speaker
                previous_speaker = speaker
                previous_timestamp = current_time
                buffer_text = text

    # Append the last entry
    if previous_speaker:
        transformed_data.append(
            (previous_speaker, previous_timestamp, previous_timestamp + timedelta(seconds=3), buffer_text.strip()))

    # Construct the formatted transcript string
    formatted_transcript = []
    for speaker, start_time, end_time, text in transformed_data:
        formatted_transcript.append(f"{start_time.hour}:{start_time.minute}:{start_time.second:.2f} --> "
                                    f"{end_time.hour}:{end_time.minute}:{end_time.second:.2f}")
        formatted_transcript.append(speaker)
        formatted_transcript.append(text)
        formatted_transcript.append("")  # Add spacing between entries

    return "\n".join(formatted_transcript)


def convertMeet(file_path):
    # Load the DOCX file
    doc = Document(file_path)

    # Extract all paragraphs as text
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Identify start and end indexes
    start_index = None
    end_index = None
    total_duration = 0  # Store total meeting duration in seconds

    for i, line in enumerate(lines):
        if line.lower() == "transcript":
            start_index = i + 1  # Start after the "Transcript" heading
        elif line.lower().startswith("meeting ended after"):
            end_index = i  # Stop before this line

            # Extract total meeting duration
            match = re.search(r"(\d+):(\d+):?(\d+)?", line)  # Matches hh:mm:ss or mm:ss formats
            if match:
                time_parts = [int(x) if x else 0 for x in match.groups()]
                if len(time_parts) == 3:  # hh:mm:ss format
                    total_duration = time_parts[0] * 3600 + time_parts[1] * 60 + time_parts[2]
                else:  # mm:ss format
                    total_duration = time_parts[0] * 60 + time_parts[1]
            break

    # If start or end markers are missing, return an empty string
    if start_index is None or end_index is None or total_duration == 0:
        return ""

    # Extract relevant transcript text
    transcript_lines = lines[start_index:end_index]

    # Process transcript data
    transformed_data = []
    current_speaker = None
    current_text = []

    # Compute total character count for proportional timestamp allocation
    total_chars = sum(len(line) for line in transcript_lines)

    if total_chars == 0:  # Avoid division by zero
        return ""

    accumulated_time = 0  # Track current timestamp in seconds

    for line in transcript_lines:
        if ":" in line:
            # A new speaker entry detected
            if current_speaker and current_text:
                duration = (sum(len(txt) for txt in current_text) / total_chars) * total_duration
                transformed_data.append(
                    (current_speaker, accumulated_time, accumulated_time + duration, " ".join(current_text)))
                accumulated_time += duration  # Update timestamp

            # Extract speaker name and text
            parts = line.split(":", 1)
            current_speaker = parts[0].strip()
            current_text = [parts[1].strip()]
        else:
            # Continuation of the same speaker's text
            current_text.append(line)

    # Append the last processed entry
    if current_speaker and current_text:
        duration = (sum(len(txt) for txt in current_text) / total_chars) * total_duration
        transformed_data.append(
            (current_speaker, accumulated_time, accumulated_time + duration, " ".join(current_text)))

    # Construct the formatted transcript string
    formatted_transcript = []
    for speaker, start_time, end_time, text in transformed_data:
        start_td = timedelta(seconds=start_time)
        end_td = timedelta(seconds=end_time)

        formatted_transcript.append(
            f"{start_td.seconds // 3600}:{(start_td.seconds % 3600) // 60}:{start_td.seconds % 60:.2f} --> "
            f"{end_td.seconds // 3600}:{(end_td.seconds % 3600) // 60}:{end_td.seconds % 60:.2f}")
        formatted_transcript.append(speaker)
        formatted_transcript.append(text)
        formatted_transcript.append("")  # Add spacing between entries

    return "\n".join(formatted_transcript)

def convertDiarization(file_path):
    # Regular expression pattern to extract timestamps and speaker
    pattern = re.compile(r"(\d+\.\d+) --> (\d+\.\d+)\n(SPEAKER_\d+)\n(.*)")

    # Read the transcript file
    with open(file_path, "r", encoding="utf-8") as file:
        content = file.read().strip()

    # Extract matches from the transcript
    matches = pattern.findall(content)

    formatted_transcript = []

    for start_time, end_time, speaker, text in matches:
        # Convert timestamps to hours, minutes, and seconds
        start_td = timedelta(seconds=float(start_time))
        end_td = timedelta(seconds=float(end_time))

        formatted_transcript.append(
            f"{start_td.seconds // 3600}:{(start_td.seconds % 3600) // 60}:{start_td.seconds % 60:.2f} --> "
            f"{end_td.seconds // 3600}:{(end_td.seconds % 3600) // 60}:{end_td.seconds % 60:.2f}")
        formatted_transcript.append(speaker)
        formatted_transcript.append(text)
        formatted_transcript.append("")  # Add spacing between entries

    return "\n".join(formatted_transcript)

def detectTranscriptFormat(file_path):
    # Check file extension
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

        # Format 1: transcript01.docx
        if re.search(r"\d+:\d+:\d+\.\d+\s-->\s\d+:\d+:\d+\.\d+", text):
            return 1

        # Format 2: meet_test_meeting.docx
        if "Transcript" in text and "Meeting ended after" in text:
            return 2

    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read().strip()

        # Format 3: zoom_test_meeting.txt
        if re.search(r"\[.+?\] \d{2}:\d{2}:\d{2}", text):
            return 3

        # Format 4: diarization_transcript.txt
        if re.search(r"\d+\.\d+\s-->\s\d+\.\d+\nSPEAKER_\d+", text):
            return 4

    # If no format is matched, return 0
    return 0

class Date:
    def __init__(self, weekday, month, day, hour, minute, ampm, year):
        match weekday:
            case "Sun":
                self.weekday = "Sunday"
            case "Mon":
                self.weekday = "Monday"
            case "Tues":
                self.weekday = "Tuesday"
            case "Wed":
                self.weekday = "Wednesday"
            case "Thu":
                self.weekday = "Thursday"
            case "Fri":
                self.weekday = "Friday"
            case "Sat":
                self.weekday = "Saturday"

        match month:
            case "Jan":
                self.month = "January"
            case "Feb":
                self.month = "February"
            case "Mar":
                self.month = "March"
            case "Apr":
                self.month = "April"
            case "May":
                self.month = "May"
            case "Jun":
                self.month = "June"
            case "Jul":
                self.month = "July"
            case "Aug":
                self.month = "August"
            case "Sep":
                self.month = "September"
            case "Oct":
                self.month = "October"
            case "Nov":
                self.month = "November"
            case "Dec":
                self.month = "December"

        self.day = day

        if int(hour) > 12 and int(hour) != 24:
            self.hour = int(hour) - 12
            self.ampm = "PM"
        elif int(hour) == 12:
            self.hour = hour
            self.ampm = "PM"
        elif int(hour) == 24:
            self.hour = 12
            self.ampm = "AM"
        else:
            self.hour = hour
            self.ampm = "AM"

        if int(minute) < 10:
            self.minute = "0" + str(minute)
        else:
            self.minute = minute

        self.year = year

    def __str__(self):
        return f"{self.weekday}, {self.month} {self.day}, {self.year} at {self.hour}:{self.minute}{self.ampm}"