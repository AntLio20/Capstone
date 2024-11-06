# Run these commands first to get redac to run:
# Install spacy: pip install spacy
# Download language model: python -m spacy download en_core_web_lg
# Install sklearn: pip install scikit-learn

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.decomposition import NMF
from sklearn.cluster import KMeans
from docx import Document
import spacy
import docx2txt
import Transcript

# Load spaCy model
nlp = spacy.load("en_core_web_lg")

# Load transcript and clean
filepath = "trans.docx"
transcript = docx2txt.process(filepath)

# Extract and clean the transcript text
cleanedTranscript = Transcript.cleanTranscript(transcript, 0, 0)

# Define off-topic keywords
offTopicKeywords = ["aside", "off topic", "off the script", "off the logs", "not important", "unrelated", "readact"]

# Split into sentences
sentences = cleanedTranscript.split(". ")

# Transform each sentence into TF-IDF (Term Frequency-Inverse Document Frequency) numerical vectors, removing common english words
vectorizer = TfidfVectorizer(stop_words='english')
tfidfMatrix = vectorizer.fit_transform(sentences)

# Topic modeling with Non-negative Matrix Factorization: NMF, grouping words that frequently appear together into n topics
nmfModel = NMF(n_components=7, random_state=42)

# Calls the NMF model and returns a matrix: nmf_topics, where each row represents a sentence, each column represents the probability score for each topic
nmfTopics = nmfModel.fit_transform(tfidfMatrix)

# Clustering with Sentence Embeddings, converting each sentence into vector representation of the semantic meaning
embeddings = [nlp(sentence).vector for sentence in sentences]
kmeans = KMeans(n_clusters=20, random_state=42)
clusters = kmeans.fit_predict(embeddings)

# Initialize Document
summarizedMeetingNotes = Document()

# Strikethrough function for visualizaton
def addStrikethrough(paragraph, text):
    run = paragraph.add_run(text)
    run.font.strike = True

# Iterate through sentences, applying topic-based redaction to topics containing keywords
for i, sentence in enumerate(sentences):
    paragraph = summarizedMeetingNotes.add_paragraph()
    cluster = clusters[i]  # Get topic/cluster ID of the current sentence

    # Check if sentence contains off-topic keyword
    if any(keyword in sentence for keyword in offTopicKeywords):
        # Apply strikethrough to this sentence
        addStrikethrough(paragraph, sentence)

        # Redact the following sentences in the same topic/cluster
        j = i + 1
        while j < len(sentences) and clusters[j] == cluster:
            paragraph = summarizedMeetingNotes.add_paragraph()
            addStrikethrough(paragraph, sentences[j])
            j += 1
    else:
        # Add normal text for on-topic sentences
        paragraph.add_run(sentence)

# Save the redacted document
summarizedMeetingNotes.save("redactTest.docx")
print("Redacted transcript saved as 'redactTest.docx'")