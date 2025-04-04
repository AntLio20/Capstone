import os
import re
import docx2txt

import torch
import bitsandbytes
from docx import Document
from peft import PeftModel
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig

base_model_path = "./DeepSeek-R1-Distill-Llama-8B"
adapter_path = "./trained_deepseek_r1_8b"

# Configure 4-bit quantization
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16,
    bnb_4bit_use_double_quant=True,
)

max_memory_mapping = {
    0: "6GiB",  # GPU 0 limit
    "cpu": "30GiB"  # CPU fallback
}

device_map = "auto"

print("Loading model...")

try:
    # Load base model with quantization
    base_model = AutoModelForCausalLM.from_pretrained(
        base_model_path,
        quantization_config=bnb_config,
        device_map=device_map,
        max_memory=max_memory_mapping
    )

    # Load the LoRA adapter on top of base model
    model = PeftModel.from_pretrained(base_model, adapter_path)

except ValueError as e:
    print("\nERROR LOADING MODEL:", e)
    print("If you get a 'dispatched on CPU' error, try lowering the GPU memory in max_memory_mapping,")
    print("or force everything on CPU by setting device_map={'': 'cpu'}.")
    raise

tokenizer = AutoTokenizer.from_pretrained(base_model_path)
model.eval()
print("Model loaded successfully.\n")


# ------------------------------------------------------------------
# 2. LLM Prompt + Generation
#    (Generates a "raw" summary from a transcript)
# ------------------------------------------------------------------

def generate_raw_summary(transcript_text):
    prompt = (
        "You are an AI that converts transcripts into meeting minutes. "
        "Use the transcript below to produce a full summary with these sections:\n\n"
        "1) Date of Meeting\n"
        "2) Start Time\n"
        "3) End Time\n"
        "4) Present Members\n"
        "5) Absent Members\n"
        "6) Agenda Approval\n"
        "7) Previous Meeting Minutes Approval\n"
        "8) Summary of Last Meeting Notes and Decisions\n"
        "9) Detailed Summary of Current Meeting (key points, discussions, proposals)\n"
        "10) Actionable Items (per person)\n"
        "11) Adjournment Time\n\n"
        "If an item is NOT mentioned, explicitly write: 'No mention in the transcript'.\n"
        "Return only the final text, with all sections. Do not add extra commentary.\n\n"
        "Transcript:\n"
        f"{transcript_text}\n\n"
        "# End of transcript"
    )

    # Tokenize
    inputs = tokenizer(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=2048,
        padding="max_length",
    )

    # Place on GPU if using GPU offload. If you're on CPU only,
    # you can remove the .to("cuda") or wrap in a try-except if needed.
    if device_map != {"": "cpu"}:
        inputs = inputs.to("cuda")

    # Generate
    with torch.no_grad():
        output_tokens = model.generate(
            **inputs,
            max_new_tokens=1200,
            temperature=0.3,
            top_p=0.9,
            repetition_penalty=1.2,
            num_return_sequences=1
        )

    raw_text = tokenizer.decode(output_tokens[0], skip_special_tokens=True)
    return raw_text


def parse_summary_to_dict(raw_summary):
    data_dict = {
        "date": "No mention in the transcript.",
        "start_time": "No mention in the transcript.",
        "end_time": "No mention in the transcript.",
        "present_members": [],
        "absent_members": [],
        "agenda_approval": "No mention in the transcript.",
        "previous_minutes_approval": "No mention in the transcript.",
        "summary_of_last_meeting": "No mention in the transcript.",
        "current_meeting_summary": [],  # store as list of bullet points
        "actionable_items": {},  # store as dict: { "Name": [task1, task2], ... }
        "adjournment_time": "No mention in the transcript."
    }

    lines = raw_summary.splitlines()
    current_section = None

    # Simple approach: look for keywords in lines
    for line in lines:
        line_str = line.strip()

        # Identify which section we are in
        if re.match(r"(?i).*date of meeting.*", line_str):
            current_section = "date"
            continue
        elif re.match(r"(?i).*start time.*", line_str):
            current_section = "start_time"
            continue
        elif re.match(r"(?i).*end time.*", line_str):
            current_section = "end_time"
            continue
        elif re.match(r"(?i).*present members.*", line_str):
            current_section = "present_members"
            continue
        elif re.match(r"(?i).*absent members.*", line_str):
            current_section = "absent_members"
            continue
        elif re.match(r"(?i).*agenda approval.*", line_str):
            current_section = "agenda_approval"
            continue
        elif re.match(r"(?i).*previous meeting minutes approval.*", line_str):
            current_section = "previous_minutes_approval"
            continue
        elif re.match(r"(?i).*summary of last meeting.*", line_str):
            current_section = "summary_of_last_meeting"
            continue
        elif re.match(r"(?i).*detailed summary of current meeting.*", line_str):
            current_section = "current_meeting_summary"
            continue
        elif re.match(r"(?i).*actionable items.*", line_str):
            current_section = "actionable_items"
            continue
        elif re.match(r"(?i).*adjournment.*", line_str):
            current_section = "adjournment_time"
            continue

        # We then parse the line content based on whichever current_section is set
        if current_section == "date":
            data_dict["date"] = line_str
        elif current_section == "start_time":
            data_dict["start_time"] = line_str
        elif current_section == "end_time":
            data_dict["end_time"] = line_str
        elif current_section == "present_members":
            if line_str:
                data_dict["present_members"].append(line_str)
        elif current_section == "absent_members":
            if line_str:
                data_dict["absent_members"].append(line_str)
        elif current_section == "agenda_approval":
            data_dict["agenda_approval"] = line_str
        elif current_section == "previous_minutes_approval":
            data_dict["previous_minutes_approval"] = line_str
        elif current_section == "summary_of_last_meeting":
            data_dict["summary_of_last_meeting"] = line_str
        elif current_section == "current_meeting_summary":
            if line_str:
                data_dict["current_meeting_summary"].append(line_str)
        elif current_section == "actionable_items":
            # Rudimentary parse: "Name: Task"
            m = re.match(r"^(.*?):(.*)$", line_str)
            if m:
                name = m.group(1).strip()
                task = m.group(2).strip()
                if name not in data_dict["actionable_items"]:
                    data_dict["actionable_items"][name] = []
                if task:
                    data_dict["actionable_items"][name].append(task)
            # else, ignore or handle differently

        elif current_section == "adjournment_time":
            data_dict["adjournment_time"] = line_str

    return data_dict


def save_minutes_formatted(meeting_data, output_filename="meeting_minutes.docx"):
    doc = Document()

    #
    # (A) Document Title
    #
    doc.add_heading("Meeting Minutes", level=0)

    #
    # (B) Opening Information
    #
    doc.add_heading("Opening Information", level=1)

    # Date
    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(meeting_data.get("date", "No mention in the transcript."))

    # Start Time
    p = doc.add_paragraph()
    p.add_run("Start Time: ").bold = True
    p.add_run(meeting_data.get("start_time", "No mention in the transcript."))

    # End Time
    p = doc.add_paragraph()
    p.add_run("End Time: ").bold = True
    p.add_run(meeting_data.get("end_time", "No mention in the transcript."))

    #
    # (C) Present Members
    #
    doc.add_heading("Present Members", level=1)
    present = meeting_data.get("present_members", [])
    if present:
        for member in present:
            doc.add_paragraph(member, style='List Bullet')
    else:
        doc.add_paragraph("No mention in the transcript.")

    #
    # (D) Absent Members
    #
    doc.add_heading("Absent Members", level=1)
    absent = meeting_data.get("absent_members", [])
    if absent:
        for member in absent:
            doc.add_paragraph(member, style='List Bullet')
    else:
        doc.add_paragraph("No mention in the transcript.")

    #
    # (E) Agenda Approval
    #
    doc.add_heading("Agenda Approval", level=1)
    doc.add_paragraph(meeting_data.get("agenda_approval", "No mention in the transcript."))

    #
    # (F) Previous Meeting Minutes Approval
    #
    doc.add_heading("Previous Meeting Minutes Approval", level=1)
    doc.add_paragraph(meeting_data.get("previous_minutes_approval", "No mention in the transcript."))

    #
    # (G) Summary of Last Meeting
    #
    doc.add_heading("Summary of Last Meeting", level=1)
    doc.add_paragraph(meeting_data.get("summary_of_last_meeting", "No mention in the transcript."))

    #
    # (H) Meeting Summary (Detailed Summary)
    #
    doc.add_heading("Meeting Summary", level=1)
    summary_points = meeting_data.get("current_meeting_summary", [])
    if summary_points:
        for item in summary_points:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No mention in the transcript.")

    #
    # (I) Actionable Items
    #
    doc.add_heading("Actionable Items", level=1)
    actions = meeting_data.get("actionable_items", {})
    if actions:
        # For each person, bullet-list each task
        for assignee, tasks in actions.items():
            para = doc.add_paragraph()
            run = para.add_run(f"{assignee}:")
            run.bold = True
            if tasks:
                for task in tasks:
                    doc.add_paragraph(task, style='List Bullet')
    else:
        doc.add_paragraph("No mention in the transcript.")

    #
    # (J) Adjournment
    #
    doc.add_heading("Adjournment", level=1)
    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(meeting_data.get("date", "No mention in the transcript."))

    p = doc.add_paragraph()
    p.add_run("Start Time: ").bold = True
    p.add_run(meeting_data.get("start_time", "No mention in the transcript."))

    p = doc.add_paragraph()
    p.add_run("End Time: ").bold = True
    p.add_run(meeting_data.get("end_time", "No mention in the transcript."))

    #
    # Save final .docx
    #
    doc.save(output_filename)
    print(f"Saved meeting minutes to {output_filename}")


def generate_minutes_doc(transcript_file, output_dir="./MeetingNotes"):
    # Make sure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # 1) Read the transcript text
    transcript_text = docx2txt.process(transcript_file)
    # 2) Get LLM raw summary
    raw_summary = generate_raw_summary(transcript_text)
    print(raw_summary)
    # 3) Parse into dictionary
    data_dict = parse_summary_to_dict(raw_summary)
    # 4) Build the final DOCX
    base_name = os.path.splitext(os.path.basename(transcript_file))[0]
    output_file = os.path.join(output_dir, f"{base_name}_minutes.docx")

    save_minutes_formatted(data_dict, output_file)

    return output_file
